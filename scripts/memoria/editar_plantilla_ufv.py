"""
Edita la plantilla UFV convertida para crear la memoria TFG con secciones 1-3.
Conserva logos, estilos y formato original de la plantilla.
"""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
import re

INPUT = "info_entregas/Plantilla_TFG_convertida.docx"
OUTPUT = "docs/memoria/TFG_Mateo_Madrigal_v2.docx"

doc = Document(INPUT)

# ── Utilidades ────────────────────────────────────────────────────

def set_text(para, text):
    """Reemplaza el texto de un párrafo conservando el estilo del primer run."""
    if para.runs:
        # Guardar formato del primer run
        first_run = para.runs[0]
        # Eliminar todos los runs
        for run in para.runs:
            run.text = ""
        first_run.text = text
    else:
        para.text = text


def add_para_after(doc, ref_para, text, style_name):
    """Añade un párrafo nuevo después de ref_para con el estilo indicado."""
    new_p = deepcopy(ref_para._element)
    # Limpiar el contenido copiado
    for child in list(new_p):
        if child.tag.endswith('}r'):
            new_p.remove(child)
    # Crear nuevo run con el texto
    from docx.oxml.ns import qn
    new_r = new_p.makeelement(qn('w:r'), {})
    # Copiar rPr del primer run del ref_para si existe
    if ref_para.runs:
        old_rPr = ref_para.runs[0]._element.find(qn('w:rPr'))
        if old_rPr is not None:
            new_r.append(deepcopy(old_rPr))
    new_t = new_r.makeelement(qn('w:t'), {})
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    # Insertar después del párrafo de referencia
    ref_para._element.addnext(new_p)
    return new_p


def clear_and_set(para, text):
    """Limpia completamente un párrafo y pone texto nuevo."""
    for run in para.runs:
        run._element.getparent().remove(run._element)
    from docx.oxml.ns import qn
    new_r = para._element.makeelement(qn('w:r'), {})
    # Obtener propiedades de run del estilo
    rPr_elem = para._element.find(qn('w:pPr'))
    if rPr_elem is not None:
        rPr_run = rPr_elem.find(qn('w:rPr'))
        if rPr_run is not None:
            new_r.append(deepcopy(rPr_run))
    new_t = new_r.makeelement(qn('w:t'), {})
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    para._element.append(new_r)


def insert_paragraphs_after(anchor_para, texts_and_styles):
    """Inserta múltiples párrafos después de anchor_para.
    texts_and_styles: lista de (text, style_para_index) donde style_para_index
    es el índice del párrafo cuyo estilo se copiará."""
    current = anchor_para._element
    for text, style_idx in reversed(texts_and_styles):
        ref = doc.paragraphs[style_idx]
        new_p = deepcopy(ref._element)
        # Limpiar runs existentes
        from docx.oxml.ns import qn
        for child in list(new_p):
            if child.tag.endswith('}r'):
                new_p.remove(child)
        # Crear run con texto
        new_r = new_p.makeelement(qn('w:r'), {})
        if ref.runs:
            old_rPr = ref.runs[0]._element.find(qn('w:rPr'))
            if old_rPr is not None:
                new_r.append(deepcopy(old_rPr))
        new_t = new_r.makeelement(qn('w:t'), {})
        new_t.text = text
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        current.addnext(new_p)


# ── Índice de estilos de referencia ───────────────────────────────
# Encontrar párrafos con cada estilo UFV para usarlos como plantilla

style_refs = {}
for i, p in enumerate(doc.paragraphs):
    sn = p.style.name if p.style else ""
    if sn not in style_refs and sn.startswith("UFV_"):
        style_refs[sn] = i

print("Estilos encontrados:")
for sn, idx in style_refs.items():
    print(f"  {sn} -> párrafo [{idx}]")

# Índices clave de la plantilla:
# [10-11] Título "PLANTILLA PARA LA ELABORACIÓN..."
# [13] Template english
# [19] Alumno:
# [20] Email:
# [21] Grado en
# [22] Curso académico
# [28] Resumen
# [30] Texto del resumen placeholder
# [32] Palabras clave placeholder
# [34] Abstract
# [36] Abstract text placeholder
# [38] Keywords placeholder
# [40] Agradecimientos
# [44] A mis padres...
# [60-120] Instrucciones de formato (eliminar)
# [122] OBJETIVOS
# [123-124] Placeholder objetivos
# [126] Introducción
# [127] Placeholder intro
# [128] Motivación
# [129] Placeholder motivación
# [130] Estado del arte
# [131] Placeholder estado del arte
# [132] Marco teórico
# [133] Placeholder marco teórico
# [134] TRABAJO TÉCNICO
# [135] Herramientas empleadas
# [136] Placeholder herramientas
# [137] Ingeniería del dato
# [138] Placeholder ingeniería
# [139] Análisis del dato
# [140] Placeholder análisis
# [141] Análisis del Negocio
# [142] Placeholder negocio
# [143] CONCLUSIONES
# [144] Placeholder conclusiones
# [145] REFERENCIAS
# [146-147] Texto explicativo refs
# [148-152] Referencias ejemplo

# ══════════════════════════════════════════════════════════════════
# PASO 1: EDITAR PORTADA
# ══════════════════════════════════════════════════════════════════

print("\n--- Editando portada ---")

# Título
set_text(doc.paragraphs[10], "Optimización de Carteras de Inversión")
set_text(doc.paragraphs[11], "mediante Predicción de Retornos con XGBoost y Señales Macroeconómicas, de Riesgo y de Sentimiento Aplicado a ETFs")

# Eliminar línea en inglés
set_text(doc.paragraphs[13], "")

# Datos del autor
set_text(doc.paragraphs[19], "Alumno: Mateo Madrigal Arteaga")
set_text(doc.paragraphs[20], "Email: mateo.madrigal@ufv.es")
set_text(doc.paragraphs[21], "Grado en Análisis de Negocios / Business Analytics")
set_text(doc.paragraphs[22], "Curso académico 2025-2026")

print("  Portada actualizada")

# ══════════════════════════════════════════════════════════════════
# PASO 2: EDITAR RESUMEN/ABSTRACT
# ══════════════════════════════════════════════════════════════════

print("--- Editando resumen/abstract ---")

set_text(doc.paragraphs[30],
    "Este trabajo presenta un sistema de optimización dinámica de carteras de inversión "
    "que combina predicciones de retornos semanales generadas por modelos de gradient "
    "boosting (XGBoost y LightGBM) con la optimización de Markowitz de media-varianza. "
    "El sistema se aplica a un universo de 10 ETFs representativos de 5 clases de activos, "
    "utilizando 109 variables predictivas agrupadas en 6 dimensiones: mercado, riesgo, "
    "macroeconómicas, liquidez, sentimiento y NLP de noticias. La validación se realiza "
    "mediante walk-forward expanding window sobre 778 semanas (2011-2026), garantizando "
    "la ausencia de data leakage. Los resultados muestran que la cartera XGBoost optimizada "
    "alcanza un Sharpe ratio de 1,397 y un retorno total acumulado del 426%, superando "
    "significativamente al benchmark 60/40 (Sharpe 0,847) y a la optimización clásica de "
    "Markowitz (Sharpe 0,832). El análisis SHAP revela que las condiciones financieras "
    "(NFCI) son el principal driver predictivo."
)

set_text(doc.paragraphs[32],
    "Palabras clave: optimización de carteras; machine learning; XGBoost; predicción de retornos; SHAP; ETFs"
)

set_text(doc.paragraphs[36],
    "This thesis presents a dynamic portfolio optimization system that combines weekly "
    "return predictions generated by gradient boosting models (XGBoost and LightGBM) with "
    "Markowitz mean-variance optimization. The system is applied to a universe of 10 ETFs "
    "representing 5 asset classes, using 109 predictive features grouped into 6 dimensions. "
    "Validation is performed through walk-forward expanding window over 778 weeks (2011-2026), "
    "ensuring no data leakage. Results show that the optimized XGBoost portfolio achieves a "
    "Sharpe ratio of 1.397 and a cumulative total return of 426%, significantly outperforming "
    "the 60/40 benchmark (Sharpe 0.847) and classic Markowitz optimization (Sharpe 0.832). "
    "SHAP analysis reveals that financial conditions (NFCI) are the primary predictive driver."
)

set_text(doc.paragraphs[38],
    "Keywords: portfolio optimization; machine learning; XGBoost; return prediction; SHAP; ETFs"
)

print("  Resumen/Abstract actualizados")

# ══════════════════════════════════════════════════════════════════
# PASO 3: EDITAR AGRADECIMIENTOS
# ══════════════════════════════════════════════════════════════════

print("--- Editando agradecimientos ---")

set_text(doc.paragraphs[44],
    "Quiero expresar mi agradecimiento a mi tutor por su orientación y seguimiento "
    "durante la elaboración de este trabajo. Agradezco también a la Universidad Francisco "
    "de Vitoria y al equipo docente del Grado en Business Analytics por la formación "
    "recibida. En la elaboración de este trabajo se ha utilizado Claude (Anthropic, versiones "
    "Claude 3.5 Sonnet y Claude Opus 4) como herramienta de asistencia para la "
    "generación y revisión de código Python. Todas las decisiones metodológicas, la "
    "interpretación de resultados y las conclusiones son responsabilidad exclusiva del autor."
)

print("  Agradecimientos actualizados")

# ══════════════════════════════════════════════════════════════════
# PASO 4: ELIMINAR INSTRUCCIONES DE FORMATO (párrafos 51-121)
# Reemplazar con contenido vacío para no romper la estructura XML
# ══════════════════════════════════════════════════════════════════

print("--- Limpiando instrucciones de formato ---")

# Párrafos 51-121 son instrucciones de la plantilla — los vaciamos
# (Nota: los índices de figuras y tablas placeholder los dejamos vacíos)
for i in range(51, 122):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        # Solo vaciar si no es un epígrafe
        if "Epígrafe" not in (p.style.name if p.style else ""):
            for run in p.runs:
                run.text = ""

# Poner contenido útil en los índices
set_text(doc.paragraphs[55], "Índice de Figuras")
set_text(doc.paragraphs[56], "[Se generará automáticamente al finalizar el documento]")
set_text(doc.paragraphs[58], "Índice de Tablas")
set_text(doc.paragraphs[59], "[Se generará automáticamente al finalizar el documento]")

print("  Instrucciones eliminadas")

# ══════════════════════════════════════════════════════════════════
# PASO 5: EDITAR SECCIONES DE CONTENIDO
# ══════════════════════════════════════════════════════════════════

print("--- Editando secciones de contenido ---")

# OBJETIVOS [122-124]
set_text(doc.paragraphs[122], "1. OBJETIVOS")
set_text(doc.paragraphs[123],
    "Objetivo general: Diseñar, implementar y evaluar un sistema de optimización dinámica "
    "de carteras de inversión basado en predicciones de retornos semanales generadas por "
    "modelos de gradient boosting (XGBoost y LightGBM), combinados con un optimizador de "
    "Markowitz de media-varianza, validado out-of-sample mediante walk-forward expanding window."
)
set_text(doc.paragraphs[124],
    "Objetivos específicos: "
    "OE1: Construir un dataset multidimensional de 109 variables predictivas a partir de 6 fuentes "
    "de datos (Yahoo Finance, FRED, Google Trends, AAII, Refinitiv). "
    "OE2: Entrenar y optimizar modelos XGBoost y LightGBM con Optuna (75 trials) para predicción "
    "de retornos semanales de 10 ETFs. "
    "OE3: Diseñar un pipeline de walk-forward con expanding window y embargo temporal. "
    "OE4: Comparar las carteras ML contra tres benchmarks (60/40, equal-weight, Markowitz) "
    "mediante Sharpe, Sortino, Max Drawdown y Calmar. "
    "OE5: Interpretar las predicciones mediante SHAP para identificar variables y dimensiones "
    "con mayor poder predictivo. "
    "OE6: Investigar la robustez mediante feature selection y enriquecimiento del dataset."
)

# INTRODUCCIÓN [126-127]
set_text(doc.paragraphs[126], "Introducción")
set_text(doc.paragraphs[127],
    "La gestión de carteras de inversión constituye uno de los problemas fundamentales "
    "de las finanzas modernas. Desde la publicación del trabajo seminal de Markowitz (1952) "
    "sobre la selección óptima de carteras, la industria financiera ha adoptado el marco "
    "de media-varianza como referencia para la asignación de activos. Sin embargo, este "
    "enfoque clásico descansa sobre supuestos que la evidencia empírica ha cuestionado "
    "repetidamente: la normalidad de los retornos, la estabilidad de las correlaciones "
    "entre activos y la estacionariedad de los parámetros estadísticos (DeMiguel et al., 2009). "
    "La crisis financiera global de 2008, la pandemia de COVID-19 en 2020 y el shock "
    "inflacionario de 2022 pusieron de manifiesto las limitaciones de los modelos estáticos. "
    "En paralelo, el campo del Machine Learning ha experimentado avances significativos "
    "en su aplicación a problemas financieros. Gu, Kelly y Xiu (2020), en su influyente "
    "estudio \"Empirical Asset Pricing via Machine Learning\", demostraron que los modelos "
    "de gradient boosting superan a los modelos lineales tradicionales en la predicción "
    "de retornos. Este trabajo se sitúa en la intersección de estos dos campos: utiliza "
    "modelos de ML para predecir retornos semanales de ETFs y alimenta estas predicciones "
    "a un optimizador de Markowitz que asigna pesos dinámicamente cada semana. La pregunta "
    "de investigación es: ¿puede un sistema basado en XGBoost y LightGBM generar carteras "
    "de ETFs que superen consistentemente a los benchmarks clásicos? El universo de inversión "
    "comprende 10 ETFs que cubren renta variable (SPY, QQQ, IWM, EFA, EEM), renta fija "
    "(AGG, LQD, TIP), oro (GLD) e inmobiliario (VNQ), con 987 semanas de datos (2007-2026) "
    "y 109 features predictivas en 6 dimensiones. El código fuente completo está disponible "
    "en GitHub: https://github.com/mateomadrigal/TFG-OPTIMIZATION-ML-1"
)

# MOTIVACIÓN [128-129]
set_text(doc.paragraphs[128], "Motivación")
set_text(doc.paragraphs[129],
    "DeMiguel, Garlappi y Uppal (2009) demostraron que una cartera equiponderada simple "
    "(1/N) supera consistentemente a la optimización de Markowitz fuera de muestra, debido "
    "a la inestabilidad de las estimaciones de covarianza. Este resultado, conocido como la "
    "\"maldición de Markowitz\", motivó la búsqueda de enfoques alternativos. Los ETFs han "
    "democratizado el acceso a la diversificación entre clases de activos, permitiendo a "
    "inversores de cualquier tamaño construir carteras globales con costes reducidos. La "
    "hipótesis central de este trabajo es que las predicciones de ML, al incorporar señales "
    "macroeconómicas, de riesgo y de sentimiento más allá de los retornos históricos, pueden "
    "generar carteras que superen consistentemente a los benchmarks clásicos. El reto principal "
    "reside en evitar el sobreajuste en un entorno donde la señal predictiva es inherentemente "
    "débil: una precisión direccional del 55-60% ya constituye un resultado notable (Gu et al., 2020)."
)

# ESTADO DEL ARTE [130-131]
set_text(doc.paragraphs[130], "Estado del arte")
set_text(doc.paragraphs[131],
    "Teoría moderna de carteras: La teoría tiene su origen en Markowitz (1952), quien formalizó "
    "la optimización media-varianza. Sharpe (1964) extendió este marco con el CAPM. Black y "
    "Litterman (1992) propusieron un modelo que combina el equilibrio de mercado con views "
    "subjetivas del gestor. "
    "Machine Learning en finanzas: Gu, Kelly y Xiu (2020) compararon modelos ML en la predicción "
    "de retornos, concluyendo que gradient boosting y redes neuronales superan significativamente "
    "a los modelos lineales. Fischer y Krauss (2018) aplicaron redes LSTM al S&P 500. De Prado "
    "(2018) enfatiza la necesidad de walk-forward validation con purging y embargo temporal. "
    "Gradient Boosting: Friedman (2001) formalizó las gradient boosting machines. Chen y Guestrin "
    "(2016) propusieron XGBoost con regularización L1/L2 y manejo nativo de NaN. Ke et al. (2017) "
    "desarrollaron LightGBM con crecimiento leaf-wise y técnicas GOSS y EFB. "
    "Interpretabilidad: Lundberg y Lee (2017) propusieron SHAP basado en la teoría de Shapley. "
    "TreeSHAP (Lundberg et al., 2020) es un algoritmo exacto para modelos de árboles. "
    "Señales alternativas: Baker y Wurgler (2006) demostraron que el sentimiento de inversores "
    "tiene poder predictivo. Da, Engelberg y Gao (2015) propusieron el índice FEARS basado en "
    "Google Trends. "
    "Optimización ML: Ban, El Karoui y Lim (2018) propusieron regularización basada en rendimiento. "
    "Yan (2025) comparó XGBoost y LightGBM en gestión de riesgo de carteras. Tashman (2000) "
    "formalizó el concepto de walk-forward validation."
)

# MARCO TEÓRICO [132-133]
set_text(doc.paragraphs[132], "Marco teórico")
set_text(doc.paragraphs[133],
    "[Se incluirá en la sección 3.3 Análisis del dato, según las indicaciones del taller "
    "de análisis del dato 2025-2026. El marco teórico detallará: la formulación matemática "
    "de XGBoost y LightGBM, la optimización de Markowitz con restricciones, la metodología "
    "walk-forward y el framework SHAP.]"
)

# HERRAMIENTAS [135-136]
set_text(doc.paragraphs[135], "Herramientas empleadas")
set_text(doc.paragraphs[136],
    "El proyecto se ha desarrollado íntegramente en Python 3.12 sobre un entorno Linux (WSL2). "
    "Las principales librerías utilizadas son: XGBoost 3.2 y LightGBM 4.6 para los modelos de "
    "gradient boosting; Optuna 4.3 para la optimización bayesiana de hiperparámetros; SHAP 0.51 "
    "para la interpretabilidad; pandas 2.2 y NumPy 2.4 para el procesamiento de datos; "
    "matplotlib 3.10 para visualización; scikit-learn 1.8 para métricas y preprocesamiento; "
    "scipy 1.17 para la optimización de Markowitz; yfinance, fredapi y pytrends para la "
    "extracción de datos. Como herramienta de asistencia se utilizó Claude (Anthropic), citado "
    "según la taxonomía CRediT. El código fuente completo y los datos procesados están "
    "disponibles en el repositorio: https://github.com/mateomadrigal/TFG-OPTIMIZATION-ML-1"
)

# INGENIERÍA DEL DATO [137-138]
set_text(doc.paragraphs[137], "Ingeniería del dato")
set_text(doc.paragraphs[138],
    "[POR INSERTAR — Contenido completo de la sección de ingeniería del dato con el proceso "
    "ETL, limpieza, feature engineering y análisis exploratorio. Ver documento "
    "ingenieria_del_dato_memoria.docx]"
)

# ANÁLISIS DEL DATO [139-140]
set_text(doc.paragraphs[139], "Análisis del dato")
set_text(doc.paragraphs[140],
    "[POR INSERTAR — Marco teórico (XGBoost, LightGBM, Markowitz, walk-forward, SHAP), "
    "entrenamiento de modelos, optimización con Optuna, evaluación de resultados, "
    "interpretación SHAP, comparativa de carteras.]"
)

# ANÁLISIS DEL NEGOCIO [141-142]
set_text(doc.paragraphs[141], "Análisis del Negocio")
set_text(doc.paragraphs[142], "[POR INSERTAR]")

# CONCLUSIONES [143-144]
set_text(doc.paragraphs[143], "CONCLUSIONES GENERALES DEL TRABAJO")
set_text(doc.paragraphs[144], "[POR INSERTAR — Conclusiones, limitaciones y líneas de trabajo futuro.]")

# ══════════════════════════════════════════════════════════════════
# PASO 6: REEMPLAZAR REFERENCIAS
# ══════════════════════════════════════════════════════════════════

print("--- Editando referencias ---")

set_text(doc.paragraphs[145], "REFERENCIAS BIBLIOGRÁFICAS")
set_text(doc.paragraphs[146], "")
set_text(doc.paragraphs[147], "")

refs = [
    "Baker, M., & Wurgler, J. (2006). Investor sentiment and the cross-section of stock returns. The Journal of Finance, 61(4), 1645-1680.",
    "Ban, G. Y., El Karoui, N., & Lim, A. E. (2018). Machine learning and portfolio optimization. Management Science, 64(3), 1136-1154.",
    "Black, F., & Litterman, R. (1992). Global portfolio optimization. Financial Analysts Journal, 48(5), 28-43.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD, 785-794.",
    "Da, Z., Engelberg, J., & Gao, P. (2015). The sum of all FEARS investor sentiment and asset prices. The Review of Financial Studies, 28(1), 1-32.",
]

# Reemplazar las 5 referencias de ejemplo con las primeras 5 nuestras
for i, ref in enumerate(refs):
    if 148 + i < len(doc.paragraphs):
        set_text(doc.paragraphs[148 + i], ref)

# Las demás las añadimos después de la última referencia existente
extra_refs = [
    "de Prado, M. L. (2018). Advances in Financial Machine Learning. John Wiley & Sons.",
    "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus naive diversification. The Review of Financial Studies, 22(5), 1915-1953.",
    "Fischer, T., & Krauss, C. (2018). Deep learning with LSTM networks for financial market predictions. European Journal of Operational Research, 270(2), 654-669.",
    "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189-1232.",
    "Gu, S., Kelly, B., & Xiu, D. (2020). Empirical asset pricing via machine learning. The Review of Financial Studies, 33(5), 2223-2273.",
    "Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30.",
    "Leippold, M., Wang, Q., & Zhou, W. (2022). Machine learning in the Chinese stock market. Journal of Financial Economics, 145(2), 64-82.",
    "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in NeurIPS, 30.",
    "Lundberg, S. M., et al. (2020). From local explanations to global understanding with explainable AI for trees. Nature Machine Intelligence, 2(1), 56-67.",
    "Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77-91.",
    "Sharpe, W. F. (1964). Capital asset prices: A theory of market equilibrium. The Journal of Finance, 19(3), 425-442.",
    "Tashman, L. J. (2000). Out-of-sample tests of forecasting accuracy. International Journal of Forecasting, 16(4), 437-450.",
    "Yan, J. (2025). Portfolio risk management with gradient boosting: LightGBM vs XGBoost. Journal of Risk and Financial Management, 18(2), 85.",
]

# Añadir después del último párrafo de referencia existente
last_ref_para = doc.paragraphs[152]
for ref in reversed(extra_refs):
    add_para_after(doc, last_ref_para, ref, "UFV_Referencias")

# Limpiar la nota sobre generación automática
if 154 < len(doc.paragraphs):
    set_text(doc.paragraphs[154], "")

print("  Referencias actualizadas (18 total)")

# ══════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
import os
print(f"\nDocumento guardado: {OUTPUT}")
print(f"Tamaño: {os.path.getsize(OUTPUT):,} bytes")
