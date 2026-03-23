"""
Inserta el contenido de ingenieria_del_dato_memoria.docx en TFG_Mateo_Madrigal_v3.docx,
reemplazando el marcador [POR INSERTAR] de la sección Ingeniería del dato.
"""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
from lxml import etree
import os

INPUT_MEMORIA = "docs/memoria/TFG_Mateo_Madrigal_v3.docx"
INPUT_INGENIERIA = "info_entregas/ingenieria_del_dato_memoria.docx"
OUTPUT = "docs/memoria/TFG_Mateo_Madrigal_v4.docx"

# Namespaces XML de docx
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
}

# ── Cargar documentos ─────────────────────────────────────────────

doc_main = Document(INPUT_MEMORIA)
doc_ing = Document(INPUT_INGENIERIA)

print(f"Memoria: {len(doc_main.paragraphs)} párrafos")
print(f"Ingeniería: {len(doc_ing.paragraphs)} párrafos, {len(doc_ing.tables)} tablas")

# ── Encontrar el marcador a reemplazar ────────────────────────────

marker_idx = None
for i, p in enumerate(doc_main.paragraphs):
    if "POR INSERTAR" in p.text and "ingeniería" in p.text.lower():
        marker_idx = i
        break

if marker_idx is None:
    print("ERROR: No se encontró marcador de ingeniería del dato")
    exit(1)

print(f"Marcador encontrado en párrafo [{marker_idx}]: {doc_main.paragraphs[marker_idx].text[:80]}")

# ── Encontrar estilos de referencia en la memoria ─────────────────

# Mapear estilos de ingeniería → estilos UFV de la memoria
style_map = {}
for p in doc_main.paragraphs:
    sn = p.style.name if p.style else ""
    if sn and sn not in style_map:
        style_map[sn] = p.style

print(f"Estilos disponibles en memoria: {list(style_map.keys())[:10]}...")

# ── Preparar el contenido a insertar ──────────────────────────────
# Saltamos los párrafos 0-16 de ingeniería (portada + título "3.2 INGENIERÍA")
# porque el heading ya existe en la memoria

# Párrafos a copiar: desde [18] (3.2.1 Origen) hasta el final
START_PARA = 18  # "3.2.1 Origen de los datos"

# Recopilar elementos a insertar (párrafos y tablas en orden)
elements_to_insert = []

# Necesitamos iterar el body XML del documento de ingeniería para mantener el orden
# de párrafos, tablas e imágenes
body_ing = doc_ing.element.body

# Contador de figuras para marcadores
fig_counter = 0
fig_para_indices = {31, 39, 47, 62, 67, 72, 77, 82, 87}  # Párrafos con imágenes

# ── Copiar relaciones de imágenes ─────────────────────────────────
# Necesitamos copiar las imágenes del doc de ingeniería al doc principal

from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Mapear rIds del doc de ingeniería a nuevos rIds en el doc principal
rid_map = {}
for rel_id, rel in doc_ing.part.rels.items():
    if 'image' in rel.reltype:
        # Copiar la imagen al doc principal
        image_part = rel.target_part
        new_rid = doc_main.part.relate_to(image_part, rel.reltype)
        rid_map[rel_id] = new_rid

print(f"Imágenes copiadas: {len(rid_map)}")

# ── Insertar elementos después del marcador ───────────────────────

marker_element = doc_main.paragraphs[marker_idx]._element

# Primero, limpiar el texto del marcador
for run in doc_main.paragraphs[marker_idx].runs:
    run.text = ""

# Recorrer los elementos del body de ingeniería en orden
para_count = 0
inserted_count = 0
current_insert_point = marker_element

for child in body_ing:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

    if tag == 'p':
        para_count += 1
        if para_count <= START_PARA:
            continue  # Saltar portada y título principal

        # Copiar el párrafo
        new_p = deepcopy(child)

        # Actualizar rIds de imágenes en el párrafo copiado
        for blip in new_p.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed and embed in rid_map:
                blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                         rid_map[embed])

        # Remapear estilos: Heading 2 → buscar equivalente UFV
        pPr = new_p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                old_style = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                # Mapear estilos
                if old_style == 'Heading2':
                    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'UFVEpgrafeNivel2')
                elif old_style == 'Heading3':
                    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'UFVEpgrafeNivel3')
                elif old_style == 'Normal':
                    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'UFVNormal')

        current_insert_point.addnext(new_p)
        current_insert_point = new_p
        inserted_count += 1

    elif tag == 'tbl':
        para_count_check = sum(1 for c in body_ing if c.tag.split('}')[-1] == 'p' and c is not child)
        # Copiar la tabla
        new_tbl = deepcopy(child)
        current_insert_point.addnext(new_tbl)
        current_insert_point = new_tbl
        inserted_count += 1

print(f"Elementos insertados: {inserted_count}")

# ── Guardar ───────────────────────────────────────────────────────

doc_main.save(OUTPUT)
size = os.path.getsize(OUTPUT)
print(f"\nGuardado: {OUTPUT} ({size:,} bytes)")

# ── Verificar ─────────────────────────────────────────────────────

doc_check = Document(OUTPUT)
print(f"\nVerificación: {len(doc_check.paragraphs)} párrafos, {len(doc_check.tables)} tablas")

print("\n=== ESTRUCTURA DE INGENIERÍA INSERTADA ===")
in_section = False
for i, p in enumerate(doc_check.paragraphs):
    sn = p.style.name if p.style else "None"
    text = p.text.strip()
    if "Ingeniería del dato" in text and "Epígrafe" in sn:
        in_section = True
    if in_section and text:
        is_heading = "Epígrafe" in sn or "Epgrafe" in sn or "Heading" in sn
        if is_heading or "3.2" in text[:10] or "Figura" in text[:10] or "POR INSERTAR" in text:
            print(f"[{i:3d}] ({sn:35s}) {text[:100]}")
    if in_section and "Análisis del dato" in text and "Epígrafe" in sn:
        break
