"""
Inserta la sección 5 (Análisis del dato) en TFG_Mateo_Madrigal_v4.docx,
reemplazando el marcador [POR INSERTAR] del análisis del dato.
También limpia el marcador del marco teórico en la sección 2.4.
"""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
from lxml import etree
import os

INPUT = "docs/memoria/TFG_Mateo_Madrigal_v4.docx"
OUTPUT = "docs/memoria/TFG_Mateo_Madrigal_v5.docx"

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc = Document(INPUT)
print(f"Documento cargado: {len(doc.paragraphs)} párrafos")

# ── Encontrar párrafos de referencia para copiar estilos ──────────

ref_nivel2 = ref_nivel3 = ref_normal = None
for p in doc.paragraphs:
    sn = p.style.name if p.style else ""
    if sn == "UFV_Epígrafe_Nivel_2" and ref_nivel2 is None:
        ref_nivel2 = p
    elif sn == "UFV_Epígrafe_Nivel_3" and ref_nivel3 is None:
        ref_nivel3 = p
    elif sn == "UFV_Normal" and p.text.strip() and ref_normal is None:
        ref_normal = p

print(f"Estilos de referencia encontrados: N2={ref_nivel2 is not None}, N3={ref_nivel3 is not None}, Normal={ref_normal is not None}")


def make_paragraph(text, style_ref):
    """Crea un nuevo elemento XML de párrafo copiando el estilo de style_ref."""
    new_p = deepcopy(style_ref._element)
    # Limpiar runs existentes
    for child in list(new_p):
        if child.tag == f'{{{WNS}}}r':
            new_p.remove(child)
    # Crear run con texto
    new_r = etree.SubElement(new_p, f'{{{WNS}}}r')
    # Copiar propiedades de run
    if style_ref.runs:
        old_rPr = style_ref.runs[0]._element.find(f'{{{WNS}}}rPr')
        if old_rPr is not None:
            new_r.insert(0, deepcopy(old_rPr))
    new_t = etree.SubElement(new_r, f'{{{WNS}}}t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p


# ── Contenido de la Sección 5: Análisis del Dato ─────────────────

content = []  # Lista de (texto, tipo) donde tipo = "h2", "h3", "p", "fig", "table_header"

# === PRIMERA CARILLA ===
content.append(("Estrategia de modelización", "h3"))
content.append((
    "Este trabajo plantea un problema de aprendizaje supervisado de regresión sobre datos "
    "tabulares de series temporales financieras semanales. El objetivo es predecir el "
    "log-return de la semana siguiente para cada uno de los 10 ETFs del universo de "
    "inversión, utilizando 109 variables predictivas que capturan información de mercado, "
    "riesgo, macro, liquidez, sentimiento y noticias.", "p"))
content.append((
    "La estructura de los datos (109 features numéricas, sin texto ni imágenes, con "
    "dependencia temporal) condiciona la selección de modelos. Se descartaron las redes "
    "neuronales profundas (LSTM, Transformer) por requerir más datos y ser menos "
    "interpretables para datos tabulares. Se descartó la regresión lineal como baseline "
    "porque asume relaciones lineales entre variables y no captura las interacciones ni "
    "las no-linealidades del mercado financiero; en su lugar, XGBoost con hiperparámetros "
    "por defecto cumple la función de baseline más honesto, al operar sobre la misma "
    "estructura de datos con mínima configuración.", "p"))
content.append((
    "Los modelos seleccionados son XGBoost (Chen y Guestrin, 2016) y LightGBM (Ke et al., "
    "2017), los dos principales algoritmos de gradient boosting para datos tabulares. "
    "Ambos manejan nativamente valores faltantes (NaN), son invariantes a la escala de "
    "las variables, capturan automáticamente interacciones entre features y permiten "
    "regularización para controlar el sobreajuste.", "p"))
content.append((
    "La validación se realiza mediante walk-forward expanding window (Tashman, 2000) con "
    "778 splits out-of-sample, embargo temporal de 1 semana y reentrenamiento semanal. "
    "Las predicciones alimentan un optimizador de Markowitz de media-varianza con "
    "restricciones long-only y máximo 40% por activo. La interpretabilidad se garantiza "
    "mediante SHAP (Lundberg y Lee, 2017). Se comparan 5 estrategias: XGBoost Tuned, "
    "LightGBM Tuned, Markowitz clásico, 60/40 y Equal-weight.", "p"))

# === 3.3.1 MARCO TEÓRICO ===
content.append(("3.3.1 Marco teórico", "h2"))

content.append(("Optimización de Markowitz", "h3"))
content.append((
    "El marco de media-varianza de Markowitz (1952) formaliza la selección de carteras "
    "como un problema de optimización. Dado un vector de retornos esperados mu y una "
    "matriz de covarianzas Sigma, el retorno esperado de una cartera con pesos w es "
    "E(Rp) = w'mu, y su varianza es sigma_p^2 = w'Sigma*w. El objetivo es maximizar el "
    "Sharpe ratio S = (Rp - Rf) / sigma_p, donde Rf = 0 (tipo libre de riesgo).", "p"))
content.append((
    "En este trabajo, el optimizador resuelve: max_w (w'mu_pred) / sqrt(w'Sigma*w), "
    "sujeto a las restricciones: suma(wi) = 1, 0 <= wi <= 0.40 para todo i. La "
    "restricción de máximo 40% por activo evita concentraciones excesivas. Se utiliza "
    "el método SLSQP de scipy.optimize.minimize. La innovación clave es que mu_pred "
    "no es la media histórica de los retornos, sino la predicción del modelo ML para "
    "la semana siguiente, mientras que Sigma se estima con datos históricos hasta la "
    "semana actual.", "p"))

content.append(("Gradient Boosting", "h3"))
content.append((
    "Los modelos de gradient boosting (Friedman, 2001) construyen un ensemble de árboles "
    "de decisión de forma secuencial. La predicción del ensemble es y_hat = suma(fk(x)) "
    "para k=1..K, donde cada árbol fk corrige los residuos del anterior. La función "
    "objetivo a minimizar es L(theta) = suma(l(yi, y_hat_i)) + suma(Omega(fk)), donde "
    "l es la función de pérdida (MSE para regresión) y Omega es el término de "
    "regularización: Omega(f) = gamma*T + 0.5*lambda*suma(wj^2), siendo T el número de "
    "hojas y wj los pesos de las hojas. El parámetro eta (learning rate) controla la "
    "contribución de cada árbol: y_hat_t = y_hat_(t-1) + eta * ft(x).", "p"))

content.append(("XGBoost", "h3"))
content.append((
    "XGBoost (Chen y Guestrin, 2016) utiliza una aproximación de Taylor de segundo orden "
    "a la función de pérdida, empleando los gradientes gi = dl/dy_hat y los hessianos "
    "hi = d^2l/dy_hat^2 de cada observación. El peso óptimo de cada hoja j es "
    "wj* = -sum(gi) / (sum(hi) + lambda), y la ganancia de un split candidato es "
    "Gain = 0.5 * [GL^2/(HL+lambda) + GR^2/(HR+lambda) - (GL+GR)^2/(HL+HR+lambda)] - gamma, "
    "donde GL, GR, HL, HR son las sumas de gradientes y hessianos en los nodos hijo "
    "izquierdo y derecho. XGBoost incorpora regularización L1 (alpha) y L2 (lambda), "
    "manejo nativo de NaN (asigna las observaciones con valores faltantes al nodo que "
    "maximiza la ganancia) y early stopping para detener el entrenamiento cuando la "
    "métrica de validación deja de mejorar.", "p"))

content.append(("LightGBM", "h3"))
content.append((
    "LightGBM (Ke et al., 2017) introduce dos innovaciones respecto a XGBoost. Primera, "
    "el crecimiento leaf-wise: en lugar de crecer el árbol nivel por nivel (level-wise), "
    "LightGBM elige la hoja con mayor ganancia potencial y la divide, lo que produce "
    "árboles más profundos y asimétricos pero más eficientes. Segunda, las técnicas GOSS "
    "(Gradient-based One-Side Sampling), que prioriza las observaciones con gradientes "
    "grandes al muestrear datos de entrenamiento, y EFB (Exclusive Feature Bundling), que "
    "agrupa features mutuamente excluyentes para reducir la dimensionalidad efectiva. "
    "Estas optimizaciones hacen que LightGBM sea significativamente más rápido que XGBoost "
    "con rendimiento comparable.", "p"))

content.append(("Walk-forward validation", "h3"))
content.append((
    "La validación walk-forward con expanding window (Tashman, 2000; de Prado, 2018) "
    "es el protocolo estándar para evaluar modelos predictivos en series temporales "
    "financieras. En cada semana t del período de test, el modelo se entrena con todos "
    "los datos disponibles desde la semana 1 hasta la semana t-2 (expanding window), se "
    "aplica un embargo de 1 semana para evitar leakage por autocorrelación, y se predice "
    "la semana t. El entrenamiento utiliza el último 15% de los datos de entrenamiento "
    "como conjunto de validación para early stopping con 50 rondas de paciencia. Este "
    "protocolo genera 778 predicciones out-of-sample (abril 2011 a febrero 2026), "
    "asegurando que ningún dato futuro contamina las predicciones.", "p"))

content.append(("SHAP: SHapley Additive exPlanations", "h3"))
content.append((
    "SHAP (Lundberg y Lee, 2017) asigna a cada variable i una contribución phi_i a la "
    "predicción de cada observación, basándose en la teoría de juegos cooperativos de "
    "Shapley. La propiedad fundamental es la eficiencia: suma(phi_i) = f(x) - E[f(x)], "
    "es decir, las contribuciones SHAP suman exactamente la diferencia entre la predicción "
    "individual y la media del modelo. TreeSHAP (Lundberg et al., 2020) es un algoritmo "
    "exacto con complejidad O(TLD^2) para modelos de árboles, donde T es el número de "
    "árboles, L el número de hojas y D la profundidad. Se utiliza para la interpretación "
    "global (importancia media de cada variable) y local (explicación de semanas concretas "
    "como la crisis SVB de marzo 2023).", "p"))

# === 3.3.2 MODELOS PREDICTIVOS ===
content.append(("3.3.2 Modelos predictivos", "h2"))

content.append(("Pipeline de entrenamiento", "h3"))
content.append((
    "El pipeline de entrenamiento sigue una arquitectura de tres capas. La Capa 1 "
    "(infraestructura) comprende la carga de datos, el generador de splits walk-forward "
    "y el cálculo de benchmarks. La Capa 2 (entrenamiento base) entrena XGBoost y "
    "LightGBM con hiperparámetros por defecto como línea base. La Capa 3 (optimización) "
    "utiliza Optuna con el sampler TPE (Tree-structured Parzen Estimator) y MedianPruner "
    "para buscar los mejores hiperparámetros mediante 75 trials bayesianos por modelo. "
    "En cada trial, se entrena el modelo en un subconjunto reducido (3 ETFs, reentrenamiento "
    "cada 4 semanas, 195 splits) para acelerar la búsqueda, y los mejores parámetros se "
    "validan con el walk-forward completo (10 ETFs, 778 splits).", "p"))

content.append(("Modelos baseline", "h3"))
content.append((
    "Los modelos con hiperparámetros por defecto (learning_rate=0.1, max_depth=6 para "
    "XGBoost; learning_rate=0.1, num_leaves=31 para LightGBM) establecen la línea base. "
    "Con parámetros por defecto, XGBoost alcanza un Sharpe ratio de 1.154 y LightGBM "
    "un Sharpe de 1.000, superando ya al benchmark 60/40 (Sharpe 0.847). Esto confirma "
    "que los modelos de gradient boosting aportan valor predictivo incluso sin optimización.", "p"))

content.append(("Optimización con Optuna", "h3"))
content.append((
    "La optimización bayesiana con Optuna exploró el siguiente espacio de hiperparámetros: "
    "learning_rate [0.005, 0.3], max_depth [3, 10], subsample [0.5, 1.0], "
    "colsample_bytree [0.3, 1.0], reg_alpha [1e-8, 10], reg_lambda [1e-8, 10], "
    "min_child_weight [1, 50] para XGBoost; con num_leaves [15, 127] y "
    "min_child_samples [5, 50] adicionales para LightGBM. Se realizaron 75 trials para "
    "cada modelo con función objetivo: minimizar el RMSE medio de los 3 ETFs representativos "
    "(SPY, AGG, GLD) en el walk-forward reducido.", "p"))
content.append((
    "Los mejores hiperparámetros encontrados para XGBoost fueron: learning_rate=0.022, "
    "max_depth=7, subsample=0.687, colsample_bytree=0.825, reg_alpha=2.0e-7, "
    "reg_lambda=1.3e-7, min_child_weight=20. Para LightGBM: learning_rate=0.006, "
    "num_leaves=41, max_depth=9, subsample=0.581, colsample_bytree=0.967, "
    "reg_alpha=5.4e-4, reg_lambda=0.023, min_child_samples=33. La validación con Optuna "
    "v2 (100 trials adicionales con rangos ampliados y penalización anti-overfitting) "
    "confirmó que estos parámetros representan el óptimo global: el trial semilla de v1 "
    "fue el mejor de los 100 trials de v2 para XGBoost.", "p"))

# Tabla de hiperparámetros
content.append(("Tabla 3. Comparativa de hiperparámetros por defecto y optimizados.", "table_header"))
content.append(("TABLA_HIPERPARAMS", "table"))

# === 3.3.3 MÉTRICAS DE EVALUACIÓN ===
content.append(("3.3.3 Métricas de evaluación", "h2"))

content.append(("Métricas de predicción (ML)", "h3"))
content.append((
    "Las métricas de evaluación de las predicciones se seleccionaron en función del "
    "problema de regresión planteado. El RMSE (Root Mean Squared Error) es la métrica "
    "primaria de entrenamiento, que penaliza proporcionalmente los errores grandes. "
    "El MAE (Mean Absolute Error) complementa al RMSE al ser menos sensible a outliers. "
    "El R² mide la proporción de varianza explicada; valores negativos son habituales en "
    "predicción financiera y no indican un modelo inútil, sino que los retornos son "
    "extremadamente ruidosos (Gu et al., 2020). La Directional Accuracy (56.8%) es la "
    "métrica más relevante en la práctica: mide el porcentaje de semanas en que el modelo "
    "acierta la dirección del retorno (positivo o negativo). No se utiliza MAPE porque "
    "explota cuando los retornos son cercanos a cero, ni AUC-ROC porque es una métrica "
    "de clasificación.", "p"))

content.append(("Métricas financieras", "h3"))
content.append((
    "Las métricas financieras evalúan el rendimiento de las carteras resultantes. "
    "El Sharpe ratio (Sharpe, 1964) es la métrica principal: S = (Rp - Rf) / sigma_p, "
    "donde Rp es el retorno anualizado y sigma_p la volatilidad anualizada, con Rf=0. "
    "El Sortino ratio reemplaza sigma_p por la volatilidad a la baja (solo retornos "
    "negativos), penalizando exclusivamente el riesgo de pérdida. El Maximum Drawdown "
    "mide la mayor caída porcentual desde un pico hasta un valle en la curva de equity. "
    "El Calmar ratio es el cociente retorno/max_drawdown. El VaR 95% semanal (-2.11%) "
    "indica que en el 5% peor de las semanas la cartera XGBoost pierde al menos ese "
    "porcentaje, y el CVaR 95% (-3.48%) es la media de esas pérdidas extremas.", "p"))

content.append(("Tabla 4. Comparativa de las 7 estrategias (walk-forward out-of-sample, 778 semanas).", "table_header"))
content.append(("TABLA_COMPARATIVA", "table"))

content.append(("Diagnóstico de overfitting", "h3"))
content.append((
    "El diagnóstico de sobreajuste se realizó comparando la desviación estándar de las "
    "predicciones con la de los retornos reales. El ratio medio es 0.16x (las predicciones "
    "son 6 veces menos volátiles que la realidad), confirmando que el modelo es conservador "
    "y no exagera. El ratio train_RMSE/test_RMSE medio es 0.92x, indicando ausencia de "
    "sobreajuste: el modelo no memoriza los datos de entrenamiento. Adicionalmente, la "
    "evolución temporal muestra que la cartera ML supera al 60/40 en 12 de los 16 años "
    "evaluados (75%), incluyendo años de crisis (2008 no se evalúa por ser período de "
    "entrenamiento, pero 2020 COVID y 2022 tipos sí).", "p"))

# === 3.3.4 VISUALIZACIÓN DE RESULTADOS ===
content.append(("3.3.4 Visualización de resultados", "h2"))

figures = [
    ("Figura 10. Frontera eficiente y carteras del TFG (2011-2026). Los puntos grises representan 5.000 carteras aleatorias. Las estrellas rojas (XGBoost Tuned, Sharpe 1.397) y naranjas (LightGBM Tuned, Sharpe 1.313) superan la frontera eficiente estática gracias al rebalanceo dinámico semanal.",
     "docs/figures/comparativa_carteras_riesgo_retorno.png"),
    ("Figura 11. Evolución de 100 euros invertidos en cada cartera (2011-2026), escala logarítmica. XGBoost Tuned: 100 euros a 526 euros (+426%). Las zonas sombreadas indican periodos de crisis (COVID-19, subidas de tipos 2022).",
     "docs/figures/equity_curves.png"),
    ("Figura 12. Top 20 variables por importancia SHAP global (media de los 10 ETFs). Las Condiciones Financieras (NFCI) dominan con el 16.5% de la importancia total. Colores por dimensión: azul=mercado, rojo=riesgo, naranja=sentimiento.",
     "docs/figures/shap_global_importance.png"),
    ("Figura 13. Contribución de cada dimensión al poder predictivo. Riesgo aporta 5 veces más información por variable (5.3%/feat) que Mercado (1.1%/feat). NLP de noticias contribuye 0%.",
     "docs/figures/shap_dimension_importance.png"),
    ("Figura 14. SHAP beeswarm para SPY (S&P 500). Dir. Accuracy: 59.6%. El NFCI alto (rojo) empuja las predicciones a la baja. El VIX alto tiene efecto negativo.",
     "docs/figures/shap_beeswarm_spy.png"),
    ("Figura 15. SHAP beeswarm para AGG (Bonos US). Dir. Accuracy: 56.9%. Escala ~100 veces menor que SPY — renta fija menos volátil.",
     "docs/figures/shap_beeswarm_agg.png"),
    ("Figura 16. SHAP beeswarm para GLD (Oro). Dir. Accuracy: 54.9%. Google:Inflación es el driver principal — el oro actúa como cobertura inflacionaria.",
     "docs/figures/shap_beeswarm_gld.png"),
    ("Figura 17. Importancia SHAP por variable y ETF (top 15). Las líneas blancas separan renta variable, renta fija y alternativos.",
     "docs/figures/shap_etf_comparison.png"),
    ("Figura 18. Evolución temporal de importancia SHAP en tres períodos. 2016-2020 incluye crisis COVID-19. Las Condiciones Financieras (NFCI) ganan importancia en periodos de crisis.",
     "docs/figures/shap_temporal_importance.png"),
    ("Figura 19. SHAP waterfall para SPY, semana del 17 de marzo de 2023 (crisis bancaria SVB). El deterioro de las condiciones financieras empuja la predicción a la baja.",
     "docs/figures/shap_waterfall_extreme.png"),
]

for fig_caption, fig_path in figures:
    content.append((f"[Insertar imagen: {fig_path}]", "p"))
    content.append((fig_caption, "fig"))

# === 3.3.5 INTERPRETACIÓN DE RESULTADOS ===
content.append(("3.3.5 Interpretación de resultados", "h2"))

content.append(("Interpretación SHAP", "h3"))
content.append((
    "El análisis SHAP revela que las tres variables más importantes a nivel global son: "
    "las Condiciones Financieras (NFCI, 16.5% de importancia), el Drawdown de Emergentes "
    "(3.7%) y el Nivel del VIX (3.5%). La dimensión de Mercado concentra el 68.4% de la "
    "importancia total (60 features), pero la dimensión de Riesgo es 5 veces más eficiente: "
    "con solo 4 features aporta el 21.3%, lo que equivale a 5.3% por variable frente al "
    "1.1% por variable de Mercado. Las 22 features de NLP de noticias contribuyen 0% al "
    "poder predictivo, coherente con que tienen 94.3% de valores faltantes.", "p"))
content.append((
    "Cada ETF tiene drivers diferentes. El NFCI domina para los ETFs de renta variable "
    "(SPY, QQQ, IWM, EFA). El Drawdown propio domina para Emergentes (EEM). Para el Oro "
    "(GLD), Google:Inflación es la variable más importante, reflejando el papel del oro "
    "como cobertura inflacionaria. Para Bonos US (AGG), el retorno propio domina debido "
    "a la alta autocorrelación de los retornos de renta fija. El modelo se adapta a "
    "diferentes regímenes: la importancia relativa de las variables cambia entre los "
    "períodos 2011-2015, 2016-2020 (crisis COVID) y 2021-2025 (subidas de tipos).", "p"))

content.append(("Experimentos de feature engineering", "h3"))
content.append((
    "Se realizaron dos experimentos adicionales para evaluar la robustez del dataset de "
    "109 features. En el Experimento 1, se añadieron 12 nuevas variables (WEI, CCSA, "
    "STLFSI4, MOVE Index, spread 10Y-3M, VIX term structure, correlación SPY-AGG, "
    "dispersión de retornos). El Sharpe bajó de 1.397 a 1.203 con 121 features, indicando "
    "que las nuevas variables aportan más ruido que señal. En el Experimento 2, se eliminaron "
    "49 features con SHAP < 0.00005 (todas las NLP, la mayoría de sentimiento y algunas "
    "macro). El Sharpe bajó a 1.271 con 60 features, demostrando que las features "
    "aparentemente irrelevantes contribuyen colectivamente al poder predictivo.", "p"))
content.append((
    "La conclusión es que 109 features es el óptimo: ni añadir ni eliminar mejora el "
    "resultado. Este hallazgo es metodológicamente relevante porque demuestra que XGBoost "
    "maneja eficientemente features de baja importancia individual sin necesidad de feature "
    "selection previa, y que documentar lo que no funciona es tan valioso como lo que sí.", "p"))

content.append(("Valor añadido del Machine Learning", "h3"))
content.append((
    "Las carteras ML superan la frontera eficiente estática porque el rebalanceo dinámico "
    "semanal adapta los pesos a las condiciones cambiantes del mercado. La diferencia entre "
    "el Sharpe estático de la frontera eficiente (1.028) y el Sharpe walk-forward del XGBoost "
    "Tuned (1.397) cuantifica el valor añadido del ML: 0.369 puntos de Sharpe, equivalentes "
    "a un incremento del 36% en la eficiencia riesgo-retorno. La cartera XGBoost Tuned "
    "reduce el maximum drawdown de -24.21% (equal weight) a -14.83%, proporcionando "
    "protección significativa en periodos de crisis.", "p"))

content.append(("Limitaciones", "h3"))
content.append((
    "Los resultados deben interpretarse considerando las siguientes limitaciones. Primero, "
    "no se incluyen costes de transacción ni slippage, que reducirían los retornos netos "
    "del rebalanceo semanal. Segundo, el período 2011-2026 fue predominantemente alcista "
    "para la renta variable estadounidense, lo que favorece las estrategias largas. Tercero, "
    "el R² negativo indica que las predicciones puntuales son débiles; el valor del modelo "
    "reside en acertar la dirección (56.8% > 50%), no la magnitud exacta. Cuarto, el "
    "universo de 10 ETFs es limitado y los resultados podrían no generalizarse a universos "
    "más amplios o a mercados menos líquidos.", "p"))

# ── Insertar contenido en el documento ────────────────────────────

print("\n--- Insertando contenido ---")

# Encontrar el marcador
marker_idx = None
for i, p in enumerate(doc.paragraphs):
    if "POR INSERTAR" in p.text and "Marco teórico" in p.text:
        marker_idx = i
        break

if marker_idx is None:
    print("ERROR: Marcador no encontrado")
    exit(1)

print(f"Marcador en [{marker_idx}]: {doc.paragraphs[marker_idx].text[:80]}")

# Limpiar el marcador
for run in doc.paragraphs[marker_idx].runs:
    run.text = ""

# Insertar todos los párrafos después del marcador (en orden inverso para addnext)
marker_elem = doc.paragraphs[marker_idx]._element

# Construir lista de elementos XML
xml_elements = []
for text, tipo in content:
    if tipo == "h2":
        xml_elements.append(make_paragraph(text, ref_nivel2))
    elif tipo == "h3":
        xml_elements.append(make_paragraph(text, ref_nivel3))
    elif tipo == "p":
        xml_elements.append(make_paragraph(text, ref_normal))
    elif tipo == "fig":
        xml_elements.append(make_paragraph(text, ref_normal))  # Pie de figura como texto normal
    elif tipo == "table_header":
        xml_elements.append(make_paragraph(text, ref_normal))
    elif tipo == "table":
        # Crear tabla como párrafo de texto (las tablas reales se añadirán en Word)
        if text == "TABLA_HIPERPARAMS":
            table_text = (
                "Parámetro | XGB Default | XGB Tuned | LGB Default | LGB Tuned\n"
                "learning_rate | 0.100 | 0.022 | 0.100 | 0.006\n"
                "max_depth | 6 | 7 | -1 | 9\n"
                "subsample | 0.8 | 0.687 | 0.8 | 0.581\n"
                "colsample_bytree | 0.8 | 0.825 | 0.8 | 0.967\n"
                "n_estimators | 1000 | 2000 | 1000 | 2000\n"
                "min_child_weight | 1 | 20 | — | —\n"
                "num_leaves | — | — | 31 | 41\n"
                "min_child_samples | — | — | 20 | 33\n"
                "reg_alpha | 0 | 2.0e-7 | 0 | 5.4e-4\n"
                "reg_lambda | 1 | 1.3e-7 | 1 | 0.023"
            )
            xml_elements.append(make_paragraph(table_text, ref_normal))
        elif text == "TABLA_COMPARATIVA":
            table_text = (
                "Estrategia | Sharpe | Ret.Anual | MaxDD | Sortino | Ret.Total\n"
                "XGB Tuned | 1.397 | 11.09% | -14.83% | 1.992 | 425.9%\n"
                "LGB Tuned | 1.313 | 10.23% | -19.60% | 1.835 | 361.8%\n"
                "XGB Default | 1.154 | 8.77% | -17.44% | 1.692 | 271.5%\n"
                "LGB Default | 1.000 | 8.15% | -21.86% | 1.322 | 238.3%\n"
                "60/40 | 0.847 | 8.60% | -21.67% | 1.029 | 262.3%\n"
                "Markowitz | 0.832 | 5.39% | -19.82% | 1.031 | 124.1%\n"
                "Equal Weight | 0.649 | 7.30% | -24.21% | 0.794 | 198.1%"
            )
            xml_elements.append(make_paragraph(table_text, ref_normal))

# Insertar en orden inverso (addnext inserta después, así que el último insertado queda primero)
for elem in reversed(xml_elements):
    marker_elem.addnext(elem)

print(f"Insertados {len(xml_elements)} elementos")

# También limpiar el marcador de marco teórico en sección 2.4 [166]
for i, p in enumerate(doc.paragraphs):
    if "Se incluirá en la sección 3.3" in p.text:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = "[Marco teórico incluido en la sección 3.3 Análisis del dato]"
        print(f"Marcador de marco teórico actualizado en [{i}]")
        break

# ── Guardar ───────────────────────────────────────────────────────

doc.save(OUTPUT)
size = os.path.getsize(OUTPUT)
print(f"\nGuardado: {OUTPUT} ({size:,} bytes)")

# ── Verificar estructura ─────────────────────────────────────────

doc_check = Document(OUTPUT)
print(f"\nVerificación: {len(doc_check.paragraphs)} párrafos")

print("\n=== ESTRUCTURA SECCIÓN ANÁLISIS DEL DATO ===")
in_section = False
for i, p in enumerate(doc_check.paragraphs):
    sn = p.style.name if p.style else "None"
    text = p.text.strip()
    if "Análisis del dato" in text and "Epígrafe" in sn:
        in_section = True
    if in_section and text:
        is_heading = "Epígrafe" in sn or "Epgrafe" in sn or "Heading" in sn
        if is_heading or "3.3" in text[:10] or "Figura" in text[:10] or "Tabla" in text[:10] or "POR INSERTAR" in text:
            print(f"[{i:3d}] ({sn:35s}) {text[:100]}")
    if in_section and "Análisis del Negocio" in text:
        break
