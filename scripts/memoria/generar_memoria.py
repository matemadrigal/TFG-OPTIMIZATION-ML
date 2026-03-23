"""
Genera el documento Word de la memoria del TFG con secciones 1-3 y formato UFV.
Autor: Mateo Madrigal Arteaga
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

OUTPUT = "docs/memoria/TFG_Mateo_Madrigal.docx"

doc = Document()

# ── Configurar estilos base ──────────────────────────────────────

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.5
pf.first_line_indent = Cm(1)

# Márgenes 2.5 cm
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading_custom(text, level=1):
    """Añade encabezado con formato Times New Roman."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {
            qn("w:ascii"): "Times New Roman",
            qn("w:hAnsi"): "Times New Roman",
            qn("w:cs"): "Times New Roman",
        })
        rPr.insert(0, rFonts)
    if level == 1:
        for run in h.runs:
            run.font.size = Pt(14)
    elif level == 2:
        for run in h.runs:
            run.font.size = Pt(12)
    return h


def add_para(text, bold=False, italic=False, align=None, indent=True, size=None):
    """Añade párrafo con formato."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman",
        qn("w:hAnsi"): "Times New Roman",
    })
    rPr.insert(0, rFonts)
    run.font.size = Pt(size or 12)
    run.bold = bold
    run.italic = italic
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_ref(text):
    """Añade una referencia bibliográfica con hanging indent."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman",
        qn("w:hAnsi"): "Times New Roman",
    })
    rPr.insert(0, rFonts)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


# ══════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════

for _ in range(4):
    add_para("", indent=False)

add_para("UNIVERSIDAD FRANCISCO DE VITORIA", bold=True, align="center", indent=False, size=14)
add_para("Facultad de Derecho, Empresa y Gobierno", italic=True, align="center", indent=False, size=12)
add_para("", indent=False)
add_para("TRABAJO FIN DE GRADO", bold=True, align="center", indent=False, size=16)
add_para("", indent=False)
add_para("Grado en Análisis de Negocios / Business Analytics", align="center", indent=False, size=12)
add_para("", indent=False)

for _ in range(2):
    add_para("", indent=False)

add_para("Optimización de Carteras de Inversión mediante", bold=True, align="center", indent=False, size=14)
add_para("Predicción de Retornos con XGBoost y Señales", bold=True, align="center", indent=False, size=14)
add_para("Macroeconómicas, de Riesgo y de Sentimiento", bold=True, align="center", indent=False, size=14)
add_para("Aplicado a ETFs", bold=True, align="center", indent=False, size=14)

for _ in range(3):
    add_para("", indent=False)

add_para("Autor: Mateo Madrigal Arteaga", align="right", indent=False, size=12)
add_para("Tutor: [Tutor asignado]", align="right", indent=False, size=12)
add_para("", indent=False)
add_para("Curso académico: 2025-2026", align="center", indent=False, size=12)
add_para("Convocatoria: Ordinaria", align="center", indent=False, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# AGRADECIMIENTOS
# ══════════════════════════════════════════════════════════════════

add_heading_custom("Agradecimientos", level=1)

add_para(
    "Quiero expresar mi agradecimiento a mi tutor por su orientación y seguimiento "
    "durante la elaboración de este trabajo. Agradezco también a la Universidad Francisco "
    "de Vitoria y al equipo docente del Grado en Business Analytics por la formación "
    "recibida a lo largo de estos años, que ha sido fundamental para poder abordar un "
    "proyecto de esta naturaleza.",
    align="justify"
)

add_para(
    "En la elaboración de este trabajo se ha utilizado Claude (Anthropic, versiones "
    "Claude 3.5 Sonnet y Claude Opus 4) como herramienta de asistencia para la "
    "planificación del proyecto, generación y revisión de código Python, y apoyo en la "
    "redacción. Todas las decisiones metodológicas, la interpretación de resultados y "
    "las conclusiones son responsabilidad exclusiva del autor. El uso de esta herramienta "
    "se detalla según la taxonomía CRediT en la sección correspondiente.",
    align="justify"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# RESUMEN / ABSTRACT
# ══════════════════════════════════════════════════════════════════

add_heading_custom("Resumen", level=1)

add_para(
    "Este trabajo presenta un sistema de optimización dinámica de carteras de inversión "
    "que combina predicciones de retornos semanales generadas por modelos de gradient "
    "boosting (XGBoost y LightGBM) con la optimización de Markowitz de media-varianza. "
    "El sistema se aplica a un universo de 10 ETFs representativos de 5 clases de activos, "
    "utilizando 109 variables predictivas agrupadas en 6 dimensiones: mercado, riesgo, "
    "macroeconómicas, liquidez, sentimiento y NLP de noticias. La validación se realiza "
    "mediante walk-forward expanding window sobre 778 semanas (2011-2026), garantizando "
    "la ausencia de data leakage. Los resultados muestran que la cartera XGBoost optimizada "
    "alcanza un Sharpe ratio de 1,397 y un retorno total acumulado del 426%, superando "
    "significativamente al benchmark 60/40 (Sharpe 0,847) y a la optimización clásica de "
    "Markowitz (Sharpe 0,832). El análisis SHAP revela que las condiciones financieras "
    "(NFCI) son el principal driver predictivo, con la dimensión de riesgo aportando 5 "
    "veces más información por variable que la de mercado.",
    align="justify", indent=False
)

add_para(
    "Palabras clave: optimización de carteras, machine learning, XGBoost, predicción "
    "de retornos, SHAP, ETFs, walk-forward.",
    bold=True, indent=False
)

add_para("", indent=False)

add_heading_custom("Abstract", level=1)

add_para(
    "This thesis presents a dynamic portfolio optimization system that combines weekly "
    "return predictions generated by gradient boosting models (XGBoost and LightGBM) with "
    "Markowitz mean-variance optimization. The system is applied to a universe of 10 ETFs "
    "representing 5 asset classes, using 109 predictive features grouped into 6 dimensions: "
    "market, risk, macroeconomic, liquidity, sentiment, and news NLP. Validation is "
    "performed through walk-forward expanding window over 778 weeks (2011-2026), ensuring "
    "no data leakage. Results show that the optimized XGBoost portfolio achieves a Sharpe "
    "ratio of 1.397 and a cumulative total return of 426%, significantly outperforming the "
    "60/40 benchmark (Sharpe 0.847) and classic Markowitz optimization (Sharpe 0.832). SHAP "
    "analysis reveals that financial conditions (NFCI) are the primary predictive driver, "
    "with the risk dimension providing 5 times more information per feature than the market "
    "dimension.",
    align="justify", indent=False
)

add_para(
    "Keywords: portfolio optimization, machine learning, XGBoost, return prediction, "
    "SHAP, ETFs, walk-forward.",
    bold=True, indent=False
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# ÍNDICE DE CONTENIDOS
# ══════════════════════════════════════════════════════════════════

add_heading_custom("Índice de Contenidos", level=1)

toc_entries = [
    ("Agradecimientos", ""),
    ("Resumen / Abstract", ""),
    ("Índice de Contenidos", ""),
    ("Índice de Figuras", ""),
    ("Índice de Tablas", ""),
    ("1. Introducción", ""),
    ("   1.1. Contexto y motivación", ""),
    ("   1.2. Planteamiento del problema", ""),
    ("   1.3. Alcance del proyecto", ""),
    ("   1.4. Estructura del documento", ""),
    ("2. Objetivos", ""),
    ("   2.1. Objetivo general", ""),
    ("   2.2. Objetivos específicos", ""),
    ("3. Estado del Arte", ""),
    ("   3.1. Teoría moderna de carteras", ""),
    ("   3.2. Machine Learning en predicción financiera", ""),
    ("   3.3. Gradient Boosting: XGBoost y LightGBM", ""),
    ("   3.4. Validación en series temporales financieras", ""),
    ("   3.5. Interpretabilidad: SHAP", ""),
    ("   3.6. Señales alternativas en finanzas", ""),
    ("   3.7. Optimización de carteras con Machine Learning", ""),
    ("4. Ingeniería del Dato", ""),
    ("5. Análisis del Dato", ""),
    ("6. Análisis del Negocio", ""),
    ("7. Conclusiones", ""),
    ("Referencias Bibliográficas", ""),
    ("Anexos", ""),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(entry)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if not entry.startswith("   "):
        run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECCIÓN 1: INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════

add_heading_custom("1. Introducción", level=1)

add_heading_custom("1.1. Contexto y motivación", level=2)

add_para(
    "La gestión de carteras de inversión constituye uno de los problemas fundamentales "
    "de las finanzas modernas. Desde la publicación del trabajo seminal de Markowitz (1952) "
    "sobre la selección óptima de carteras, la industria financiera ha adoptado el marco "
    "de media-varianza como referencia para la asignación de activos. Sin embargo, este "
    "enfoque clásico descansa sobre supuestos que la evidencia empírica ha cuestionado "
    "repetidamente: la normalidad de los retornos, la estabilidad de las correlaciones "
    "entre activos y la estacionariedad de los parámetros estadísticos (DeMiguel et al., 2009).",
    align="justify"
)

add_para(
    "La crisis financiera global de 2008, la pandemia de COVID-19 en 2020 y el shock "
    "inflacionario de 2022 pusieron de manifiesto las limitaciones de los modelos "
    "estáticos: en periodos de crisis, las correlaciones entre clases de activos se "
    "disparan, la diversificación deja de funcionar y las carteras optimizadas con datos "
    "históricos sufren pérdidas severas. DeMiguel et al. (2009) demostraron que una "
    "cartera equiponderada simple (1/N) supera consistentemente a la optimización de "
    "Markowitz fuera de muestra, debido a la inestabilidad de las estimaciones de "
    "covarianza.",
    align="justify"
)

add_para(
    "En paralelo, el campo del Machine Learning ha experimentado avances significativos "
    "en su aplicación a problemas financieros. Gu, Kelly y Xiu (2020), en su influyente "
    "estudio \"Empirical Asset Pricing via Machine Learning\", demostraron que los modelos "
    "de gradient boosting, y en particular XGBoost, superan a los modelos lineales "
    "tradicionales en la predicción de retornos de activos, capturando relaciones no "
    "lineales e interacciones entre variables que los modelos clásicos ignoran.",
    align="justify"
)

add_para(
    "Los ETFs (Exchange-Traded Funds) han democratizado el acceso a la diversificación "
    "entre clases de activos, permitiendo a inversores de cualquier tamaño construir "
    "carteras globales con costes reducidos. Un universo de 10 ETFs bien seleccionados "
    "puede cubrir renta variable de distintas geografías, renta fija, materias primas e "
    "inmobiliario, proporcionando los ingredientes necesarios para una optimización "
    "multi-activo efectiva.",
    align="justify"
)

add_para(
    "Este trabajo se sitúa en la intersección de estos dos campos: utiliza modelos de "
    "Machine Learning (XGBoost y LightGBM) para predecir retornos semanales de ETFs, "
    "y alimenta estas predicciones a un optimizador de Markowitz que asigna pesos "
    "dinámicamente cada semana. La hipótesis central es que las predicciones de ML, "
    "al incorporar señales macroeconómicas, de riesgo y de sentimiento más allá de los "
    "retornos históricos, pueden generar carteras que superen consistentemente a los "
    "benchmarks clásicos.",
    align="justify"
)

add_heading_custom("1.2. Planteamiento del problema", level=2)

add_para(
    "La pregunta de investigación que guía este trabajo es: ¿puede un sistema basado "
    "en XGBoost y LightGBM, alimentado con señales macroeconómicas, de riesgo, de "
    "liquidez y de sentimiento, generar carteras de ETFs que superen consistentemente "
    "a los benchmarks clásicos (60/40, equal-weight, Markowitz) en métricas de "
    "rentabilidad ajustada al riesgo?",
    align="justify"
)

add_para(
    "El reto principal reside en evitar el sobreajuste (overfitting) en un entorno donde "
    "la señal predictiva es inherentemente débil. Los retornos financieros son "
    "extremadamente ruidosos: una precisión direccional del 55-60% ya constituye un "
    "resultado notable en la literatura (Gu et al., 2020). Para garantizar la validez de "
    "los resultados, se implementa un protocolo de validación walk-forward con expanding "
    "window, embargo temporal entre entrenamiento y test, y reentrenamiento semanal, "
    "siguiendo las mejores prácticas de de Prado (2018).",
    align="justify"
)

add_heading_custom("1.3. Alcance del proyecto", level=2)

add_para(
    "El alcance del proyecto se define por los siguientes elementos:",
    align="justify"
)

bullets = [
    "Universo de inversión: 10 ETFs que cubren 5 clases de activos — renta variable "
    "(SPY, QQQ, IWM, EFA, EEM), renta fija (AGG, LQD, TIP), oro (GLD) e inmobiliario (VNQ).",
    "Período de datos: marzo 2007 a febrero 2026 (987 semanas), abarcando la Gran "
    "Crisis Financiera, la pandemia de COVID-19 y el ciclo de subidas de tipos de 2022-2024.",
    "Variables predictivas: 109 features organizadas en 6 dimensiones — mercado (60), "
    "riesgo (4), macro (4), liquidez (4), sentimiento (15) y NLP de noticias (22).",
    "Validación: walk-forward expanding window con 778 splits out-of-sample, embargo "
    "de 1 semana, y reentrenamiento semanal.",
    "Benchmarks de comparación: cartera 60/40 (SPY/AGG), cartera equiponderada (1/10 "
    "en cada ETF), y optimización de Markowitz walk-forward clásica.",
    "Restricciones de inversión: posiciones solo largas (long-only), máximo 40% en un "
    "solo activo, pesos suman 100%.",
]

for bullet in bullets:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(bullet)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

add_heading_custom("1.4. Estructura del documento", level=2)

add_para(
    "El documento se organiza en las siguientes secciones. La Sección 2 presenta los "
    "objetivos generales y específicos del trabajo. La Sección 3 revisa el estado del "
    "arte en optimización de carteras, Machine Learning financiero e interpretabilidad. "
    "La Sección 4 detalla el proceso de ingeniería del dato: extracción, transformación, "
    "limpieza, feature engineering y análisis exploratorio. La Sección 5 aborda el "
    "análisis del dato: marco teórico, entrenamiento de modelos, optimización de "
    "hiperparámetros con Optuna, evaluación de resultados e interpretación con SHAP. "
    "La Sección 6 presenta el análisis de negocio con las implicaciones prácticas de los "
    "resultados. Finalmente, la Sección 7 recoge las conclusiones, limitaciones y líneas "
    "de trabajo futuro.",
    align="justify"
)

add_para(
    "El código fuente completo del proyecto está disponible en el repositorio público "
    "de GitHub: https://github.com/mateomadrigal/TFG-OPTIMIZATION-ML-1",
    align="justify", italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECCIÓN 2: OBJETIVOS
# ══════════════════════════════════════════════════════════════════

add_heading_custom("2. Objetivos", level=1)

add_heading_custom("2.1. Objetivo general", level=2)

add_para(
    "Diseñar, implementar y evaluar un sistema de optimización dinámica de carteras de "
    "inversión basado en predicciones de retornos semanales generadas por modelos de "
    "gradient boosting (XGBoost y LightGBM), combinados con un optimizador de Markowitz "
    "de media-varianza, validado out-of-sample mediante walk-forward expanding window.",
    align="justify"
)

add_heading_custom("2.2. Objetivos específicos", level=2)

objetivos = [
    ("OE1", "Construir un dataset multidimensional de 109 variables predictivas a partir "
     "de 6 fuentes de datos (Yahoo Finance, FRED, Google Trends, AAII, Refinitiv), "
     "cubriendo 5 dimensiones informativas (mercado, macro, riesgo, liquidez, sentimiento) "
     "más NLP de noticias, con un proceso ETL completamente reproducible."),
    ("OE2", "Entrenar y optimizar modelos XGBoost y LightGBM con Optuna (75 trials por "
     "modelo) para la predicción de retornos semanales de 10 ETFs, implementando early "
     "stopping y regularización para prevenir el sobreajuste."),
    ("OE3", "Diseñar un pipeline de walk-forward con expanding window, embargo temporal "
     "de 1 semana y reentrenamiento semanal que garantice la validez out-of-sample de "
     "los resultados sobre 778 semanas de test."),
    ("OE4", "Comparar las carteras ML-optimizadas contra tres benchmarks financieros "
     "clásicos (60/40, equal-weight, Markowitz walk-forward) mediante métricas de "
     "rentabilidad ajustada al riesgo: Sharpe, Sortino, Max Drawdown y Calmar."),
    ("OE5", "Interpretar las predicciones del modelo mediante SHAP (SHapley Additive "
     "exPlanations) para identificar las variables y dimensiones que aportan mayor "
     "poder predictivo, generando visualizaciones interpretables."),
    ("OE6", "Investigar la robustez del modelo mediante feature selection basada en "
     "SHAP y enriquecimiento del dataset con variables adicionales, documentando los "
     "resultados de cada experimento."),
]

for code, text in objetivos:
    add_para(f"{code}: {text}", align="justify")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECCIÓN 3: ESTADO DEL ARTE
# ══════════════════════════════════════════════════════════════════

add_heading_custom("3. Estado del Arte", level=1)

# 3.1
add_heading_custom("3.1. Teoría moderna de carteras", level=2)

add_para(
    "La teoría moderna de carteras tiene su origen en el trabajo seminal de Markowitz "
    "(1952), quien formalizó el problema de selección de carteras como una optimización "
    "de media-varianza: dado un conjunto de activos con retornos esperados y una matriz "
    "de covarianzas, el inversor busca la combinación de pesos que maximice el retorno "
    "esperado para un nivel dado de riesgo, o equivalentemente, que minimice el riesgo "
    "para un retorno objetivo. El conjunto de carteras óptimas forma la llamada frontera "
    "eficiente.",
    align="justify"
)

add_para(
    "Sharpe (1964) extendió este marco con el Capital Asset Pricing Model (CAPM), "
    "introduciendo el concepto de cartera de mercado y el ratio que lleva su nombre "
    "(Sharpe ratio = exceso de retorno / volatilidad), que se ha convertido en la métrica "
    "estándar de rentabilidad ajustada al riesgo. Posteriormente, Black y Litterman "
    "(1992) propusieron un modelo que combina el equilibrio de mercado con las views "
    "subjetivas del gestor, mitigando parcialmente la sensibilidad de Markowitz a "
    "errores en las estimaciones de retornos esperados.",
    align="justify"
)

add_para(
    "Sin embargo, la aplicación práctica de la optimización de Markowitz presenta "
    "limitaciones bien documentadas. DeMiguel, Garlappi y Uppal (2009) demostraron "
    "en un estudio con 14 modelos de asignación de activos que la cartera equiponderada "
    "1/N supera consistentemente a las carteras optimizadas fuera de muestra, debido a "
    "que los errores de estimación de la media y la covarianza amplifican las posiciones "
    "extremas. Este resultado, conocido como la \"maldición de Markowitz\", motivó la "
    "búsqueda de enfoques alternativos para la asignación de activos.",
    align="justify"
)

# 3.2
add_heading_custom("3.2. Machine Learning en predicción financiera", level=2)

add_para(
    "La aplicación de Machine Learning a la predicción de retornos financieros ha "
    "experimentado un crecimiento exponencial en la última década. El trabajo de "
    "referencia es Gu, Kelly y Xiu (2020), \"Empirical Asset Pricing via Machine "
    "Learning\", publicado en The Review of Financial Studies. Los autores compararon "
    "modelos lineales, árboles de decisión, random forests, gradient boosting y redes "
    "neuronales en la predicción de retornos mensuales de acciones estadounidenses, "
    "concluyendo que los modelos de gradient boosting y las redes neuronales superan "
    "significativamente a los modelos lineales, capturando relaciones no lineales e "
    "interacciones entre variables macroeconómicas y características de las acciones.",
    align="justify"
)

add_para(
    "Fischer y Krauss (2018) aplicaron redes LSTM (Long Short-Term Memory) a la "
    "predicción de retornos diarios de acciones del S&P 500, obteniendo una precisión "
    "direccional del 56% y retornos significativos después de costes de transacción. "
    "Leippold, Wang y Zhou (2022) extendieron el análisis de Gu et al. al contexto de "
    "gestión de carteras, demostrando que las predicciones de ML mejoran la asignación "
    "de activos cuando se combinan con restricciones de inversión realistas.",
    align="justify"
)

add_para(
    "Un reto fundamental en la predicción financiera es el sobreajuste. Los retornos "
    "de activos financieros tienen una ratio señal-ruido extremadamente baja: la mayor "
    "parte de la variación es imprevisible. De Prado (2018) enfatiza la necesidad de "
    "protocolos de validación rigurosos, incluyendo walk-forward validation con purging "
    "y embargo temporal, para evitar que los resultados de backtesting sobrestimen el "
    "rendimiento real del modelo.",
    align="justify"
)

# 3.3
add_heading_custom("3.3. Gradient Boosting: XGBoost y LightGBM", level=2)

add_para(
    "Los modelos de gradient boosting, formalizados por Friedman (2001), construyen "
    "un ensemble de árboles de decisión de forma secuencial, donde cada árbol nuevo "
    "corrige los errores residuales del anterior. Este enfoque ha demostrado un "
    "rendimiento excepcional en problemas con datos tabulares, siendo el algoritmo "
    "dominante en competiciones de ciencia de datos durante la última década.",
    align="justify"
)

add_para(
    "XGBoost (eXtreme Gradient Boosting), propuesto por Chen y Guestrin (2016), "
    "incorpora regularización L1 y L2 en la función objetivo, manejo nativo de "
    "valores faltantes (NaN) y un algoritmo de búsqueda de splits aproximado que "
    "escala eficientemente a grandes volúmenes de datos. Su capacidad para trabajar "
    "directamente con datos incompletos es particularmente relevante en finanzas, "
    "donde los indicadores macroeconómicos se publican con frecuencias heterogéneas.",
    align="justify"
)

add_para(
    "LightGBM, desarrollado por Ke et al. (2017) en Microsoft Research, introdujo "
    "dos innovaciones que mejoran la eficiencia computacional: el crecimiento del "
    "árbol leaf-wise (en lugar del level-wise de XGBoost) y las técnicas GOSS "
    "(Gradient-based One-Side Sampling) y EFB (Exclusive Feature Bundling). Estas "
    "optimizaciones permiten entrenar modelos significativamente más rápido sin "
    "sacrificar precisión, lo que resulta especialmente ventajoso cuando se entrenan "
    "miles de modelos en un walk-forward semanal.",
    align="justify"
)

# 3.4
add_heading_custom("3.4. Validación en series temporales financieras", level=2)

add_para(
    "La validación de modelos predictivos en series temporales financieras requiere "
    "protocolos específicos que respeten la estructura temporal de los datos. Tashman "
    "(2000) formalizó el concepto de walk-forward validation: en lugar de una única "
    "partición train-test, el modelo se reentrena periódicamente con una ventana "
    "de entrenamiento que se expande (expanding window) o se desplaza (sliding window), "
    "y se evalúa en el periodo inmediatamente posterior.",
    align="justify"
)

add_para(
    "De Prado (2018) introdujo los conceptos de purging (eliminación de muestras del "
    "entrenamiento que se solapan temporalmente con el test) y embargo (un gap temporal "
    "entre entrenamiento y test para evitar leakage por autocorrelación). En este "
    "trabajo se implementa un embargo de 1 semana entre el último dato de entrenamiento "
    "y la semana de test, y se reentrena el modelo cada semana con todos los datos "
    "disponibles hasta ese momento (expanding window).",
    align="justify"
)

# 3.5
add_heading_custom("3.5. Interpretabilidad: SHAP", level=2)

add_para(
    "La interpretabilidad de los modelos de Machine Learning es un requisito creciente "
    "tanto en la investigación académica como en la regulación financiera. Lundberg y "
    "Lee (2017) propusieron SHAP (SHapley Additive exPlanations), un marco unificado "
    "basado en la teoría de juegos cooperativos de Shapley que asigna a cada variable "
    "una contribución marginal a la predicción de cada observación individual.",
    align="justify"
)

add_para(
    "TreeSHAP (Lundberg et al., 2020), publicado en Nature Machine Intelligence, "
    "es un algoritmo exacto y eficiente para calcular valores SHAP en modelos de "
    "árboles, con complejidad O(TLD²) donde T es el número de árboles, L el número "
    "de hojas y D la profundidad. Esta eficiencia permite calcular explicaciones "
    "completas para miles de observaciones en segundos, facilitando tanto la "
    "interpretación global (qué variables son más importantes en general) como la "
    "local (por qué el modelo predijo X para una semana concreta).",
    align="justify"
)

# 3.6
add_heading_custom("3.6. Señales alternativas en finanzas", level=2)

add_para(
    "La incorporación de datos alternativos (alternative data) a los modelos de "
    "inversión es una tendencia consolidada en la industria de gestión de activos. "
    "Baker y Wurgler (2006) demostraron que el sentimiento de los inversores tiene "
    "poder predictivo sobre los retornos de las acciones, especialmente en activos "
    "más difíciles de valorar y arbitrar.",
    align="justify"
)

add_para(
    "Da, Engelberg y Gao (2015) propusieron el índice FEARS (Financial and Economic "
    "Attitudes Revealed by Search), basado en volúmenes de búsqueda de Google Trends "
    "para términos relacionados con el miedo financiero (\"recession\", \"unemployment\", "
    "\"bear market\"), demostrando que los picos en estas búsquedas predicen caídas "
    "temporales del mercado seguidas de reversiones. La encuesta de sentimiento AAII "
    "(American Association of Individual Investors) proporciona una medida directa del "
    "sentimiento retail, que históricamente ha funcionado como indicador contrarian.",
    align="justify"
)

add_para(
    "En el ámbito del procesamiento de lenguaje natural (NLP) aplicado a noticias "
    "financieras, el algoritmo VADER (Valence Aware Dictionary and sEntiment Reasoner) "
    "se ha consolidado como una herramienta estándar para el análisis de sentimiento "
    "de textos cortos como titulares, al no requerir entrenamiento previo y estar "
    "validado para el dominio financiero.",
    align="justify"
)

# 3.7
add_heading_custom("3.7. Optimización de carteras con Machine Learning", level=2)

add_para(
    "La combinación de predicciones de ML con optimización de carteras es un campo "
    "de investigación activo. Ban, El Karoui y Lim (2018) propusieron un enfoque de "
    "regularización basada en rendimiento que penaliza carteras con retornos inestables, "
    "mejorando la robustez fuera de muestra. Yan (2025) comparó XGBoost y LightGBM "
    "en el contexto específico de gestión de riesgo de carteras, concluyendo que ambos "
    "modelos superan a las técnicas tradicionales cuando se combinan con validación "
    "temporal adecuada.",
    align="justify"
)

add_para(
    "El enfoque adoptado en este trabajo sigue la línea de integrar las predicciones "
    "de retornos del ML como inputs del vector de retornos esperados (μ) del optimizador "
    "de Markowitz, reemplazando las medias históricas por predicciones dinámicas. Esta "
    "combinación preserva las ventajas del marco de media-varianza (diversificación, "
    "restricciones de inversión) mientras incorpora la capacidad del ML para capturar "
    "relaciones no lineales y cambios de régimen.",
    align="justify"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# MARCADORES PARA SECCIONES 4-7
# ══════════════════════════════════════════════════════════════════

add_heading_custom("4. Ingeniería del Dato", level=1)
add_para("[SECCIÓN POR INSERTAR — Contenido de ingenieria_del_dato_memoria.docx]",
         bold=True, italic=True, align="center", indent=False)

doc.add_page_break()

add_heading_custom("5. Análisis del Dato", level=1)
add_para("[SECCIÓN POR INSERTAR — Marco teórico, modelos, métricas, SHAP, resultados]",
         bold=True, italic=True, align="center", indent=False)

doc.add_page_break()

add_heading_custom("6. Análisis del Negocio", level=1)
add_para("[SECCIÓN POR INSERTAR]",
         bold=True, italic=True, align="center", indent=False)

doc.add_page_break()

add_heading_custom("7. Conclusiones", level=1)
add_para("[SECCIÓN POR INSERTAR]",
         bold=True, italic=True, align="center", indent=False)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# REFERENCIAS BIBLIOGRÁFICAS
# ══════════════════════════════════════════════════════════════════

add_heading_custom("Referencias Bibliográficas", level=1)

refs = [
    "Baker, M., & Wurgler, J. (2006). Investor sentiment and the cross-section of stock returns. The Journal of Finance, 61(4), 1645-1680. https://doi.org/10.1111/j.1540-6261.2006.00885.x",
    "Ban, G. Y., El Karoui, N., & Lim, A. E. (2018). Machine learning and portfolio optimization. Management Science, 64(3), 1136-1154. https://doi.org/10.1287/mnsc.2016.2644",
    "Black, F., & Litterman, R. (1992). Global portfolio optimization. Financial Analysts Journal, 48(5), 28-43. https://doi.org/10.2469/faj.v48.n5.28",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794. https://doi.org/10.1145/2939672.2939785",
    "Da, Z., Engelberg, J., & Gao, P. (2015). The sum of all FEARS investor sentiment and asset prices. The Review of Financial Studies, 28(1), 1-32. https://doi.org/10.1093/rfs/hhu072",
    "de Prado, M. L. (2018). Advances in Financial Machine Learning. John Wiley & Sons.",
    "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus naive diversification: How inefficient is the 1/N portfolio strategy? The Review of Financial Studies, 22(5), 1915-1953. https://doi.org/10.1093/rfs/hhm075",
    "Fischer, T., & Krauss, C. (2018). Deep learning with long short-term memory networks for financial market predictions. European Journal of Operational Research, 270(2), 654-669. https://doi.org/10.1016/j.ejor.2017.11.054",
    "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. The Annals of Statistics, 29(5), 1189-1232. https://doi.org/10.1214/aos/1013203451",
    "Gu, S., Kelly, B., & Xiu, D. (2020). Empirical asset pricing via machine learning. The Review of Financial Studies, 33(5), 2223-2273. https://doi.org/10.1093/rfs/hhaa009",
    "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30, 3146-3154.",
    "Leippold, M., Wang, Q., & Zhou, W. (2022). Machine learning in the Chinese stock market. Journal of Financial Economics, 145(2), 64-82. https://doi.org/10.1016/j.jfineco.2021.08.017",
    "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765-4774.",
    "Lundberg, S. M., Erion, G., Chen, H., DeGrave, A., Prutkin, J. M., Nair, B., Katz, R., Himmelfarb, J., Bansal, N., & Lee, S. I. (2020). From local explanations to global understanding with explainable AI for trees. Nature Machine Intelligence, 2(1), 56-67. https://doi.org/10.1038/s42256-019-0138-9",
    "Markowitz, H. (1952). Portfolio selection. The Journal of Finance, 7(1), 77-91. https://doi.org/10.2307/2975974",
    "Sharpe, W. F. (1964). Capital asset prices: A theory of market equilibrium under conditions of risk. The Journal of Finance, 19(3), 425-442. https://doi.org/10.2307/2977928",
    "Tashman, L. J. (2000). Out-of-sample tests of forecasting accuracy: An analysis and review. International Journal of Forecasting, 16(4), 437-450. https://doi.org/10.1016/S0169-2070(00)00065-0",
    "Yan, J. (2025). Portfolio risk management with gradient boosting: A comparative study of LightGBM and XGBoost. Journal of Risk and Financial Management, 18(2), 85.",
]

for ref in refs:
    add_ref(ref)

# ── Guardar ──────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"Documento guardado: {OUTPUT}")
print(f"Tamaño: {os.path.getsize(OUTPUT):,} bytes")
