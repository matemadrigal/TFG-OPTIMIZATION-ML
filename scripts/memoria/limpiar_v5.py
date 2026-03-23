"""
Limpia basura de plantilla y convierte tablas pipe en tablas Word reales.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from copy import deepcopy
from lxml import etree
import os, re

INPUT = "docs/memoria/TFG_Mateo_Madrigal_v5.docx"
OUTPUT = "docs/memoria/TFG_Mateo_Madrigal_v6.docx"
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc = Document(INPUT)
body = doc.element.body
print(f"Cargado: {len(doc.paragraphs)} párrafos, {len(doc.tables)} tablas")

# ══════════════════════════════════════════════════════════════════
# 1. ELIMINAR TABLA DE EJEMPLO (Grupos/Curso 1º)
# ══════════════════════════════════════════════════════════════════

removed_tables = 0
for child in list(body):
    if child.tag.endswith('}tbl'):
        # Buscar texto en la tabla
        texts = [t.text for t in child.iter(qn('w:t')) if t.text]
        all_text = ' '.join(texts)
        if 'Grupos' in all_text and 'Curso' in all_text and 'Tomada' in all_text:
            body.remove(child)
            removed_tables += 1
            print("  Eliminada tabla de ejemplo (Grupos/Curso)")

# ══════════════════════════════════════════════════════════════════
# 2. ELIMINAR PÁRRAFOS DE BASURA DE PLANTILLA
# ══════════════════════════════════════════════════════════════════

trash_patterns = [
    'Manual de Publicaciones de la American',
    'Nota: La lista de referencias se genera',
    'Nota:La lista de referencias',
    'Logotipo de la Universidad Francisco',
    'insertar título',
    'Encabezado_Tabla',  # "Tabla 1. Número de alumnos"
    'Número de alumnos por curso',
]

# Figuras de ejemplo de la plantilla (APA citation examples)
figure_trash = [
    'Figura 1. Logotipo',
    'Figura 2. Citación de referencias',
    'Figura 3. Estilos básicos de citación',
    'Figura 4. Cita textual corta',
    'Figura 5. Cita textual corta',
    'Figura 6. Cita textual larga',
]

removed_paras = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue

    should_remove = False

    # Basura por texto
    for pattern in trash_patterns:
        if pattern in text:
            should_remove = True
            break

    # Figuras de ejemplo de la plantilla
    for fig in figure_trash:
        if fig in text:
            should_remove = True
            break

    # Encabezado de tabla de ejemplo
    if text == 'Tabla 1. Número de alumnos por curso':
        should_remove = True

    if should_remove:
        # Eliminar el elemento XML del body
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            removed_paras += 1

print(f"  Eliminados {removed_paras} párrafos de basura")

# ══════════════════════════════════════════════════════════════════
# 3. REDUCIR PÁRRAFOS VACÍOS CONSECUTIVOS (máximo 2)
# ══════════════════════════════════════════════════════════════════

# Recargar párrafos después de eliminaciones
empty_removed = 0
prev_empty = 0
for child in list(body):
    if child.tag.endswith('}p'):
        # Verificar si está vacío
        texts = [t.text for t in child.iter(qn('w:t')) if t.text and t.text.strip()]
        has_drawing = child.findall('.//' + qn('w:drawing')) or child.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')

        if not texts and not has_drawing:
            prev_empty += 1
            if prev_empty > 2:
                body.remove(child)
                empty_removed += 1
        else:
            prev_empty = 0

print(f"  Eliminados {empty_removed} párrafos vacíos excesivos")

# ══════════════════════════════════════════════════════════════════
# 4. CONVERTIR TABLAS PIPE EN TABLAS WORD REALES
# ══════════════════════════════════════════════════════════════════

# Recargar doc para tener párrafos actualizados
doc.save('/tmp/temp_clean.docx')
doc = Document('/tmp/temp_clean.docx')
body = doc.element.body

tables_converted = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or '|' not in text:
        continue

    lines = text.split('\n')
    if len(lines) < 3 or '|' not in lines[0]:
        continue

    # Parsear tabla
    rows_data = []
    for line in lines:
        line = line.strip()
        if '|' in line:
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]  # Eliminar vacíos
            rows_data.append(cells)

    if len(rows_data) < 2:
        continue

    n_cols = len(rows_data[0])

    # Verificar que todas las filas tienen el mismo número de columnas
    rows_data = [r for r in rows_data if len(r) == n_cols]
    if len(rows_data) < 2:
        continue

    print(f"  Convirtiendo tabla pipe [{i}]: {len(rows_data)} filas x {n_cols} columnas")

    # Crear tabla Word
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # El estilo no existe en esta plantilla, se usa el default

    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = cell_text
            # Formato: primera fila en negrita
            for par in cell.paragraphs:
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.space_before = Pt(2)
                for run in par.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True

    # Insertar la tabla ANTES del párrafo pipe y eliminar el párrafo
    p._element.addprevious(table._tbl)
    parent = p._element.getparent()
    if parent is not None:
        parent.remove(p._element)

    tables_converted += 1

print(f"  Convertidas {tables_converted} tablas pipe → Word")

# ══════════════════════════════════════════════════════════════════
# 5. GUARDAR
# ══════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
size = os.path.getsize(OUTPUT)
print(f"\nGuardado: {OUTPUT} ({size:,} bytes)")

# Limpiar temp
os.remove('/tmp/temp_clean.docx')

# ══════════════════════════════════════════════════════════════════
# 6. VERIFICAR
# ══════════════════════════════════════════════════════════════════

doc_check = Document(OUTPUT)
print(f"\nVerificación: {len(doc_check.paragraphs)} párrafos, {len(doc_check.tables)} tablas")

# Verificar que no queda basura
trash_found = False
for i, p in enumerate(doc_check.paragraphs):
    t = p.text.strip()
    for kw in ['Manual de Publicaciones', 'Tomada de Autor', 'Nota: La lista',
               'Nota:La lista', 'Logotipo de la Universidad']:
        if kw in t:
            print(f"  ERROR: Basura en [{i}]: {t[:80]}")
            trash_found = True

# Verificar que no quedan tablas pipe
for i, p in enumerate(doc_check.paragraphs):
    t = p.text.strip()
    if '|' in t and len(t) > 50 and '\n' in t:
        print(f"  AVISO: Posible tabla pipe sin convertir en [{i}]: {t[:80]}")

if not trash_found:
    print("  OK: No se encontró basura de plantilla")

# Mostrar estructura de secciones
print("\n=== ESTRUCTURA FINAL DE SECCIONES ===")
for i, p in enumerate(doc_check.paragraphs):
    sn = p.style.name if p.style else "None"
    t = p.text.strip()
    if t and ("Epígrafe" in sn or "Epgrafe" in sn):
        print(f"[{i:3d}] ({sn:35s}) {t[:100]}")

print(f"\n=== TABLAS ===")
for ti, table in enumerate(doc_check.tables):
    first = table.rows[0].cells[0].text.strip() if table.rows else ''
    print(f"Tabla {ti}: {len(table.rows)}x{len(table.columns)} | \"{first[:40]}\"")
