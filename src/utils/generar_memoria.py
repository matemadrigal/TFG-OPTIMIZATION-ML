"""
Genera el documento Word de la entrega intermedia (Ingeniería del Dato).
Formato UFV: Times New Roman, APA, márgenes 3/2.5 cm.
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Rutas ──────────────────────────────────────────────────────

DIR_TEMPLATES = os.path.join("docs", "templates")
DIR_FIGURES = os.path.join("docs", "figures")
OUTPUT = os.path.join("docs", "ingenieria_del_dato_memoria.docx")

LOGO_UFV = os.path.join(DIR_TEMPLATES, "image1.jpg")
LOGO_SAS = os.path.join(DIR_TEMPLATES, "image2.jpg")

fig_counter = [0]


# ── Utilidades XML ─────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Colorea el fondo de una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _get_or_create_tblPr(table):
    """Obtiene o crea el tblPr de una tabla."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    return tblPr


def remove_table_borders(table):
    """Elimina TODOS los bordes de una tabla (invisible)."""
    tblPr = _get_or_create_tblPr(table)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def remove_cell_borders(cell):
    """Fuerza todos los bordes de UNA CELDA individual a none."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)


def remove_cell_margins(cell):
    """Elimina márgenes/padding internos de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="0" w:type="dxa"/>'
        '  <w:left w:w="0" w:type="dxa"/>'
        '  <w:bottom w:w="0" w:type="dxa"/>'
        '  <w:right w:w="0" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(margins)


def set_table_borders(table, color="BFBFBF", size="4"):
    """Aplica bordes finos uniformes a una tabla."""
    tblPr = _get_or_create_tblPr(table)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_page_field(paragraph):
    """Inserta campo PAGE en un párrafo."""
    r1 = paragraph.add_run()
    r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = paragraph.add_run()
    r2._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    ))
    r3 = paragraph.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


# ── Funciones de formato ───────────────────────────────────────

def add_paragraph(doc, text, font_size=11, bold=False, italic=False,
                   alignment=None, space_before=0, space_after=6,
                   color=None, first_indent=True):
    """Párrafo body: Times New Roman, interlineado 1.5, sangría 1.27 cm."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if alignment is not None:
        pf.alignment = alignment
    # Sangría APA primera línea 1.27 cm (media pulgada)
    if first_indent and alignment is None:
        pf.first_line_indent = Cm(1.27)
    return p


def add_heading_styled(doc, text, level=1):
    """
    Heading estilo UFV:
      H1 → Times New Roman 14pt, negrita, MAYÚSCULAS
      H2 → Times New Roman 12pt, negrita
      H3 → Times New Roman 11pt, negrita cursiva
    """
    display = text.upper() if level == 1 else text
    h = doc.add_heading(display, level=level)

    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(14)
            run.font.italic = False
        elif level == 2:
            run.font.size = Pt(12)
            run.font.italic = False
        else:
            run.font.size = Pt(11)
            run.font.italic = True

    pf = h.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    if level == 1:
        pf.space_before = Pt(24)
        pf.space_after = Pt(12)
    elif level == 2:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
    else:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
    return h


def add_figure(doc, img_path, caption, width_cm=13):
    """Figura centrada + pie cursiva 10pt. Sin espacio imagen→pie, sí después."""
    fig_counter[0] += 1

    # Imagen centrada
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(0)
    p_img.paragraph_format.line_spacing = 1.0
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    # Pie de figura: cursiva 10pt centrado, interlineado sencillo
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after = Pt(0)
    p_cap.paragraph_format.line_spacing = 1.0

    run_num = p_cap.add_run(f"Figura {fig_counter[0]}. ")
    run_num.font.name = "Times New Roman"
    run_num.font.size = Pt(10)
    run_num.font.italic = True
    run_num.font.bold = True

    run_txt = p_cap.add_run(caption)
    run_txt.font.name = "Times New Roman"
    run_txt.font.size = Pt(10)
    run_txt.font.italic = True

    # Párrafo vacío DESPUÉS del pie (separar del texto siguiente)
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after = Pt(6)
    p_sep.paragraph_format.line_spacing = 1.0


def add_table(doc, headers, rows, col_widths=None):
    """Tabla: cabecera #1F3864, alternas #F2F2F2, bordes #BFBFBF, 10pt."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Ancho 100% del área de texto (5000 = 100% en fiftieths of a percent)
    tblPr = _get_or_create_tblPr(table)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))

    # Bordes finos grises en toda la tabla
    set_table_borders(table, color="BFBFBF", size="4")

    # Cabecera azul oscuro con texto blanco bold
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = hdr
        set_cell_shading(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas de datos con alternas grises
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            if i % 2 == 0:
                set_cell_shading(cell, "F2F2F2")

    # Ancho de columnas
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)

    # Espacio después de la tabla
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    return table


# ── Configuración del documento ───────────────────────────────

def _make_invisible_table(table):
    """Hace una tabla COMPLETAMENTE invisible: sin estilo, sin bordes, sin margins."""
    # Quitar estilo de tabla (evita bordes heredados)
    tblPr = _get_or_create_tblPr(table)
    for old in tblPr.findall(qn("w:tblStyle")):
        tblPr.remove(old)

    # Bordes de tabla a none
    remove_table_borders(table)

    # Cell margins de tabla a 0
    for old in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(old)
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '  <w:top w:w="0" w:type="dxa"/>'
        '  <w:left w:w="0" w:type="dxa"/>'
        '  <w:bottom w:w="0" w:type="dxa"/>'
        '  <w:right w:w="0" w:type="dxa"/>'
        '</w:tblCellMar>'
    ))

    # Bordes y margins de CADA CELDA individual a none/0
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)
            remove_cell_margins(cell)


def _add_logo_table_to_header(header_obj):
    """Crea tabla 1×2 invisible con logos UFV/SAS dentro de un header."""
    # Limpiar párrafos existentes del header
    for p in header_obj.paragraphs:
        p.text = ""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Crear tabla 1×2 dentro del header
    h_table = header_obj.add_table(rows=1, cols=2, width=Cm(15.5))
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _make_invisible_table(h_table)

    # Logo UFV izquierda
    cell_l = h_table.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.paragraph_format.space_before = Pt(0)
    p_l.paragraph_format.space_after = Pt(0)
    p_l.add_run().add_picture(LOGO_UFV, width=Cm(4.5))

    # Logo SAS derecha
    cell_r = h_table.rows[0].cells[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_before = Pt(0)
    p_r.paragraph_format.space_after = Pt(0)
    p_r.add_run().add_picture(LOGO_SAS, width=Cm(2.8))


def setup_document(doc):
    """Márgenes, logos en header (todas las páginas), footer con nº página."""
    section = doc.sections[0]

    # Márgenes
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # Primera página distinta (portada sin nº de página)
    section.different_first_page_header_footer = True

    # ── Header con logos en TODAS las páginas ──
    # Header por defecto (páginas 2+)
    header = section.header
    header.is_linked_to_previous = False
    _add_logo_table_to_header(header)

    # Header primera página (portada): también con logos
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    _add_logo_table_to_header(first_header)

    # ── Footer: portada vacío, resto con nº página ──
    # Footer primera página: vacío (sin número)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    for p in first_footer.paragraphs:
        p.text = ""

    # Footer por defecto: — N —
    footer = section.footer
    footer.is_linked_to_previous = False
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(0)
    pf.paragraph_format.space_after = Pt(0)

    r1 = pf.add_run("— ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_page_field(pf)

    r2 = pf.add_run(" —")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Numeración empieza en 0 → portada=0 (oculto), contenido=1
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    sectPr.append(parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="0"/>'))


# ── Portada ────────────────────────────────────────────────────

def crear_portada(doc):
    """Portada: solo título y datos (logos vienen del header). Salto al final."""

    # ~10 líneas vacías para centrar título verticalmente
    # (los logos ya están arriba gracias al header)
    for _ in range(10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # Título principal
    add_paragraph(doc,
        "Ingeniería del Dato — Entrega Intermedia",
        font_size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=16, color=(0x1F, 0x38, 0x64),
        first_indent=False)

    # Subtítulo
    add_paragraph(doc,
        "Optimización de Carteras de Inversión mediante Predicción de Retornos "
        "con XGBoost y Señales Macroeconómicas, de Riesgo, de Sentimiento "
        "Aplicado a ETFs",
        font_size=13, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=36, color=(0x33, 0x33, 0x33), first_indent=False)

    # Datos del alumno (sin separador unicode)
    for texto in ["Alumno: Mateo Madrigal Arteaga",
                   "Grado en Business Analytics",
                   "Universidad Francisco de Vitoria",
                   "Curso 2025-2026"]:
        add_paragraph(doc, texto,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=4, space_after=4, first_indent=False)

    # Salto de página después de la portada
    doc.add_page_break()


# ── Secciones del contenido ────────────────────────────────────

def seccion_321_origen(doc):
    """3.2.1 Origen de los datos."""
    add_heading_styled(doc, "3.2.1 Origen de los datos", level=2)

    add_paragraph(doc,
        "El dataset construido para este trabajo integra información de cinco dimensiones "
        "complementarias, seleccionadas por su relevancia teórica para la predicción de retornos "
        "de activos financieros y la optimización de carteras. Cada dimensión captura un aspecto "
        "diferente del entorno de mercado: la dimensión financiera recoge los precios históricos "
        "de los ETFs que componen el universo de inversión; la macroeconómica incorpora indicadores "
        "del ciclo económico; la dimensión de riesgo mide el estrés del sistema financiero; la de "
        "liquidez monitoriza los flujos monetarios de la Reserva Federal; y la de sentimiento "
        "captura la percepción de los inversores a través de encuestas, búsquedas en Internet y "
        "análisis de noticias financieras.")

    add_paragraph(doc,
        "Adicionalmente, se incorporó una sexta fuente de datos procedente de la plataforma "
        "Refinitiv/LSEG Workspace, consistente en 17.181 titulares de noticias financieras sobre "
        "los 10 ETFs del universo de inversión, que fueron procesados mediante técnicas de "
        "procesamiento de lenguaje natural (NLP) para extraer señales de sentimiento.")

    add_paragraph(doc,
        "La tabla 1 resume las fuentes de datos, las series extraídas, la frecuencia original "
        "de publicación y el método de extracción utilizado. Todos los datos fueron obtenidos "
        "mediante código reproducible alojado en los scripts del directorio src/extractors/ del "
        "repositorio del proyecto.",
        space_after=12)

    add_table(doc,
        headers=["Dimensión", "Fuente", "Series", "Frecuencia", "Método"],
        rows=[
            ["Financiera", "Yahoo Finance", "10 ETFs (SPY, QQQ, AGG, GLD...)",
             "Diaria", "yfinance API"],
            ["Macroeconómica", "Federal Reserve (FRED)", "8 series (CPI, UNRATE, DGS10...)",
             "Diaria/Mensual", "fredapi"],
            ["Riesgo", "Federal Reserve (FRED)", "4 series (VIX, HY spread, NFCI...)",
             "Diaria/Semanal", "fredapi"],
            ["Liquidez", "Federal Reserve (FRED)", "4 series (WALCL, RRPONTSYD...)",
             "Semanal", "fredapi"],
            ["Sentimiento", "Google Trends / AAII", "7 términos + encuesta AAII",
             "Semanal/Mensual", "pytrends / scraping"],
            ["Noticias NLP", "Refinitiv/LSEG", "17.181 titulares (10 ETFs)",
             "Intradía", "refinitiv-data API"],
        ],
        col_widths=[2.5, 3.1, 4.0, 2.5, 3.4])


def seccion_322_etl(doc):
    """3.2.2 Proceso ETL."""
    add_heading_styled(doc, "3.2.2 Proceso ETL", level=2)

    add_paragraph(doc,
        "El proceso de extracción, transformación y carga (ETL) se diseñó como un pipeline "
        "secuencial de scripts Python ejecutables independientemente, siguiendo el principio de "
        "reproducibilidad científica. Cada script lee datos de la etapa anterior, aplica "
        "transformaciones documentadas y guarda el resultado en la siguiente carpeta del pipeline "
        "(raw → interim → processed).")

    add_heading_styled(doc, "Extracción", level=3)

    add_paragraph(doc,
        "Los datos financieros de los 10 ETFs se descargaron mediante la librería yfinance, "
        "que proporciona acceso programático a los precios históricos de Yahoo Finance desde "
        "enero de 2007 hasta la fecha actual. Las series macroeconómicas, de riesgo y de liquidez "
        "se obtuvieron de la API de FRED (Federal Reserve Economic Data) mediante la librería "
        "fredapi, implementando lógica de reintentos con backoff exponencial para gestionar "
        "errores temporales del servidor. Los datos de sentimiento de Google Trends se extrajeron "
        "con la librería pytrends, segmentando las consultas en ventanas de 5 años para cumplir "
        "con las restricciones de la API y aplicando pausas entre solicitudes para evitar "
        "bloqueos por exceso de peticiones (error 429). Finalmente, los titulares de noticias "
        "se obtuvieron de Refinitiv/LSEG Workspace mediante su API oficial, implementando una "
        "arquitectura de proxy para resolver las restricciones de conexión entre el entorno "
        "WSL2 y la aplicación de escritorio en Windows.")

    add_heading_styled(doc, "Transformación: alineación temporal", level=3)

    add_paragraph(doc,
        "El principal reto de la transformación fue alinear series con frecuencias heterogéneas "
        "(diaria, semanal y mensual) a una frecuencia común. Se eligió la frecuencia semanal "
        "con cierre en viernes (W-FRI) como estándar, por tres razones: primero, el viernes es el "
        "último día de negociación bursátil de la semana; segundo, la frecuencia semanal reduce "
        "el ruido de los datos diarios sin perder información relevante sobre tendencias y crisis; "
        "tercero, es coherente con la frecuencia de publicación de la mayoría de las señales de "
        "sentimiento y liquidez.")

    add_paragraph(doc,
        "El método de resampleo aplicado dependió de la frecuencia original de cada serie. Para "
        "las series diarias (precios de ETFs, tipos de interés) se tomó el último valor observado "
        "de cada semana. Para las series mensuales (CPI, desempleo) se aplicó forward-fill antes "
        "del resampleo, reflejando que el último dato publicado sigue siendo la mejor estimación "
        "disponible hasta la siguiente publicación. Las series semanales se alinearon al calendario "
        "de viernes mediante reindexación con forward-fill.")

    add_paragraph(doc,
        "La Figura 1 muestra el efecto del resampleo sobre la serie de precios de SPY, "
        "ilustrando cómo la serie semanal preserva toda la información relevante con menor ruido.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "antes_despues_frecuencia.png"),
        "Resampleo de frecuencia diaria (n=4.826) a semanal (n=1.000) del precio de SPY. "
        "La serie semanal preserva tendencias y crisis con menor ruido intradía.",
        width_cm=13)


def seccion_323_limpieza(doc):
    """3.2.3 Limpieza y tratamiento de datos."""
    add_heading_styled(doc, "3.2.3 Limpieza y tratamiento de datos", level=2)

    add_heading_styled(doc, "Tratamiento de valores nulos", level=3)

    add_paragraph(doc,
        "Tras la alineación temporal, el dataset presentaba valores nulos distribuidos de forma "
        "desigual entre dimensiones. La dimensión de sentimiento concentraba el 98% de los nulos "
        "(4.003 de un total de 4.032), debido a dos factores: la encuesta AAII se publica los "
        "jueves y al alinear a viernes no siempre coincidía, y los datos mensuales de Google "
        "Trends generaban huecos entre meses. Las dimensiones macroeconómica y de riesgo "
        "presentaban nulos marginales (13 y 16 respectivamente).")

    add_paragraph(doc,
        "El método de imputación aplicado fue forward-fill seguido de backward-fill. Esta "
        "decisión se fundamenta en la práctica estándar del análisis financiero: el último dato "
        "publicado de un indicador económico sigue siendo la mejor estimación disponible hasta "
        "que se publica el siguiente. No se trata de inventar datos, sino de reflejar cómo los "
        "inversores trabajan con la información disponible en cada momento. El caso especial de "
        "la serie RRPONTSYD (Reverse Repo de la Fed) se rellenó con cero antes de 2013, "
        "puesto que el programa no existía antes de esa fecha y el volumen de operaciones era "
        "literalmente nulo.")

    add_paragraph(doc,
        "La Figura 2 muestra el contraste visual entre los datos de sentimiento antes y "
        "después de la limpieza, evidenciando la efectividad del tratamiento.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "antes_despues_nulos.png"),
        "Heatmap de valores nulos en la dimensión de sentimiento antes (4.003 nulos, 40%) "
        "y después (0 nulos) del tratamiento con forward-fill.",
        width_cm=13)

    add_heading_styled(doc, "Eliminación de duplicados entre dimensiones", level=3)

    add_paragraph(doc,
        "Se identificaron tres series que aparecían simultáneamente en más de una dimensión. "
        "Se decidió mantener cada serie en la dimensión donde tiene mayor relevancia conceptual: "
        "el VIX y el spread de High Yield se mantuvieron en la dimensión de riesgo (eliminándose "
        "de macro), y el RRPONTSYD se mantuvo en liquidez (eliminándose de riesgo). Tras la "
        "deduplicación, la dimensión macroeconómica pasó de 10 a 8 columnas y la de riesgo de "
        "5 a 4.")

    add_heading_styled(doc, "Análisis de outliers", level=3)

    add_paragraph(doc,
        "Se identificaron outliers en los retornos semanales de los 10 ETFs utilizando el "
        "método del rango intercuartílico (IQR), clasificando como outlier toda observación "
        "fuera del intervalo [Q1 − 1,5 × IQR, Q3 + 1,5 × IQR]. El porcentaje de outliers "
        "osciló entre el 2,6% (GLD) y el 5,9% (VNQ), con una media del 4% por ETF.")

    add_paragraph(doc,
        "Se tomó la decisión de mantener todos los outliers en el dataset. Esta decisión se "
        "fundamenta en que los valores extremos corresponden a eventos reales del mercado "
        "(crisis financiera de 2008, pandemia de COVID-19, shock inflacionario de 2022) y "
        "eliminarlos supondría privar al modelo de información sobre los escenarios más "
        "extremos, que son precisamente los que más importa predecir en la gestión de carteras.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "outliers_analisis.png"),
        "Análisis de outliers por método IQR. Panel superior: número de outliers por ETF. "
        "Panel inferior: distribución temporal de outliers de SPY, concentrados en periodos "
        "de crisis.",
        width_cm=13)

    add_heading_styled(doc, "Normalización", level=3)

    add_paragraph(doc,
        "Se aplicó normalización z-score (StandardScaler de scikit-learn) a todas las features "
        "para que variables con escalas muy diferentes (VIX: 9-80, CPI: 200-320, log-returns: "
        "±0,10) tengan el mismo peso relativo en el modelo. Los targets no se normalizan para "
        "que las predicciones sean directamente interpretables como log-returns. Es importante "
        "señalar que, en la fase de modelado, el scaler se ajustará exclusivamente con datos de "
        "entrenamiento para evitar fuga de información del futuro al pasado (data leakage).")


def seccion_324_features(doc):
    """3.2.4 Feature Engineering."""
    add_heading_styled(doc, "3.2.4 Feature Engineering", level=2)

    add_paragraph(doc,
        "Los datos crudos (precios, niveles de indicadores) por sí solos no son informativos "
        "para un modelo predictivo. Lo que aporta poder predictivo son las transformaciones "
        "que capturan dinámica, tendencia, riesgo y cambio. Se generaron 109 features organizadas "
        "en seis categorías, cada una con justificación financiera documentada en el código fuente.")

    add_table(doc,
        headers=["Categoría", "Nº", "Variables representativas", "Justificación"],
        rows=[
            ["ETFs", "60", "log-return, volatilidad 4/12 sem., momentum, drawdown",
             "Capturan rendimiento, riesgo y tendencia de cada activo"],
            ["Macro", "4", "Spread 10Y-2Y, variación CPI, UNRATE, UMCSENT",
             "Señales del ciclo económico y riesgo de recesión"],
            ["Riesgo", "4", "VIX nivel/cambio, HY spread, NFCI",
             "Indicadores de estrés y aversión al riesgo"],
            ["Liquidez", "4", "Var. balance Fed, reverse repo, depósitos, TGA",
             "Flujos monetarios que afectan precios de activos"],
            ["Sentimiento", "15", "AAII Bull-Bear, 7 Google Trends (cambio + MA4)",
             "Percepción de inversores retail y público general"],
            ["Noticias NLP", "22", "Sentimiento VADER por ETF + agregado",
             "Tono de noticias financieras (17.181 titulares)"],
        ],
        col_widths=[2.3, 1.2, 5.5, 6.5])

    add_paragraph(doc,
        "Para las 6 métricas por ETF (60 features en total), se eligieron transformaciones "
        "estándar de la literatura de finanzas cuantitativas. El log-return semanal "
        "(ln(P_t/P_{t-1})) se prefiere al retorno simple porque es aditivo en el tiempo y "
        "estabiliza la varianza. La volatilidad rolling en ventanas de 4 y 12 semanas captura "
        "el riesgo a corto y medio plazo. El momentum (retorno acumulado en 4 y 12 semanas) "
        "explota el efecto momentum ampliamente documentado en la literatura académica. El "
        "drawdown mide la caída desde el máximo histórico, proporcionando una medida de estrés "
        "del activo.",
        space_before=6)

    add_heading_styled(doc, "Análisis de sentimiento NLP", level=3)

    add_paragraph(doc,
        "Se aplicó el algoritmo VADER (Valence Aware Dictionary and sEntiment Reasoner) "
        "sobre los 17.181 titulares de noticias extraídos de Refinitiv/LSEG Workspace. VADER "
        "fue seleccionado por estar diseñado específicamente para textos cortos como titulares, "
        "no requerir entrenamiento previo (usa un lexicón validado por humanos) y estar "
        "ampliamente citado en la literatura de finanzas cuantitativas. Cada titular recibió un "
        "compound score entre −1 (muy negativo) y +1 (muy positivo), que se agregó a frecuencia "
        "semanal por ETF calculando la media del sentimiento y el número de titulares como "
        "proxy de atención mediática. La distribución resultante mostró un 61,7% de titulares "
        "neutros, un 23,6% positivos y un 14,7% negativos.")


def seccion_325_eda(doc):
    """3.2.5 Análisis Exploratorio."""
    add_heading_styled(doc, "3.2.5 Análisis Exploratorio de Datos", level=2)

    add_paragraph(doc,
        "El análisis exploratorio reveló varios hallazgos que condicionan directamente el "
        "diseño del modelo predictivo y justifican la elección de técnicas de Machine Learning "
        "frente a métodos tradicionales de optimización de carteras.")

    add_heading_styled(doc, "Evolución de precios y heterogeneidad de rendimientos", level=3)

    add_paragraph(doc,
        "La Figura 4 muestra la evolución de los 10 ETFs normalizada a base 100, en escala "
        "logarítmica para permitir una comparación justa. Se observa una dispersión enorme entre "
        "clases de activos: QQQ (Nasdaq 100) multiplicó su valor por más de 8 veces, mientras "
        "que EEM (emergentes) prácticamente no generó retorno neto en 19 años. Las tres crisis "
        "sombreadas se manifiestan de forma distinta en cada activo, lo que justifica la "
        "diversificación multi-activo.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "etfs_evolucion_log.png"),
        "Evolución de precios de 10 ETFs normalizada a base 100, escala logarítmica (2007–2026). "
        "Sombreados: GFC 2008, COVID-19 y crisis de inflación/tipos 2022.",
        width_cm=13)

    add_heading_styled(doc, "No normalidad de retornos", level=3)

    add_paragraph(doc,
        "La Figura 5 presenta los histogramas de retornos semanales con la curva normal "
        "teórica superpuesta. Todos los ETFs exhiben leptocurtosis (colas más pesadas que la "
        "normal, con curtosis de hasta 53,4 para AGG) y asimetría negativa (caídas extremas "
        "más frecuentes que subidas extremas). Este hallazgo invalida el supuesto de normalidad "
        "del modelo de Markowitz y justifica el uso de modelos de Machine Learning capaces de "
        "capturar relaciones no lineales y eventos extremos.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "distribucion_retornos.png"),
        "Distribución de retornos semanales de los 10 ETFs frente a la distribución normal "
        "teórica. Se observa leptocurtosis generalizada (colas pesadas).",
        width_cm=13)

    add_heading_styled(doc, "Volatilidad y crisis", level=3)

    add_paragraph(doc,
        "El índice VIX, conocido como el «índice del miedo», registró picos de 79,1 durante "
        "la Gran Crisis Financiera y de 66,0 durante la pandemia de COVID-19 (Figura 6). Su "
        "comportamiento asimétrico (sube de forma abrupta y baja gradualmente) y su clara "
        "relación negativa con los retornos de SPY justifican su inclusión como feature del "
        "modelo tanto en nivel como en variación semanal.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "vix_historico.png"),
        "Evolución del VIX (2007–2026) con niveles de referencia y picos de crisis anotados.",
        width_cm=13)

    add_heading_styled(doc, "Inestabilidad de correlaciones", level=3)

    add_paragraph(doc,
        "Uno de los hallazgos más relevantes para la motivación del proyecto es la inestabilidad "
        "de las correlaciones entre clases de activos. La Figura 7 muestra la correlación rolling "
        "de 52 semanas entre SPY (acciones) y AGG (bonos). La correlación oscila entre −0,6 y "
        "+0,6 a lo largo de los 19 años: en periodos normales es negativa (la diversificación "
        "funciona), pero en crisis como 2008, 2020 y especialmente 2022-2024 se vuelve positiva "
        "(acciones y bonos caen juntos, la diversificación se rompe). Este resultado invalida el "
        "supuesto fundamental del modelo de Markowitz de correlaciones estables y justifica un "
        "enfoque dinámico basado en Machine Learning que se adapte al régimen de correlaciones "
        "actual.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "rolling_corr_spy_agg.png"),
        "Correlación rolling de 52 semanas entre SPY y AGG. Las zonas rojas indican correlación "
        "positiva (diversificación rota); las verdes, correlación negativa (diversificación "
        "funcional).",
        width_cm=13)

    add_heading_styled(doc, "Complementariedad entre dimensiones", level=3)

    add_paragraph(doc,
        "La Figura 8 muestra la correlación media absoluta entre las cinco dimensiones del "
        "dataset. Liquidez y sentimiento presentan correlaciones muy bajas con el resto de "
        "dimensiones (0,05–0,09), confirmando que aportan información diferenciada y "
        "complementaria al modelo. Esto valida la decisión de diseño de incorporar múltiples "
        "fuentes de datos heterogéneas.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "correlacion_entre_dimensiones.png"),
        "Correlación media absoluta entre dimensiones del dataset. "
        "Liquidez y sentimiento aportan información independiente.",
        width_cm=9)

    add_heading_styled(doc, "Estacionariedad de las series", level=3)

    add_paragraph(doc,
        "Se aplicó el test Augmented Dickey-Fuller a 19 variables representativas (Figura 9). "
        "El resultado fue que 17 de 19 son estacionarias (p < 0,05), incluyendo todos los "
        "log-returns y variaciones semanales. Las dos excepciones (spread 10Y-2Y y spread "
        "Bull-Bear AAII) son variables de nivel con reversión lenta a la media, que se "
        "mantuvieron sin diferenciar porque su valor absoluto tiene significado económico directo "
        "y XGBoost, a diferencia de los modelos ARIMA, puede trabajar con variables no "
        "estacionarias sin problemas.",
        space_after=6)

    add_figure(doc,
        os.path.join(DIR_FIGURES, "tabla_estacionariedad.png"),
        "Resultados del test Augmented Dickey-Fuller. Verde: estacionaria (p < 0,05). "
        "Rojo: no estacionaria.",
        width_cm=12)


def seccion_326_master(doc):
    """3.2.6 Dataset maestro."""
    add_heading_styled(doc, "3.2.6 Dataset maestro", level=2)

    add_paragraph(doc,
        "El resultado del proceso de ingeniería del dato es un dataset maestro unificado con "
        "987 semanas de observaciones y 119 columnas, compuesto por 109 features predictoras "
        "y 10 variables target (una por ETF). El periodo cubierto abarca desde marzo de 2007 "
        "hasta febrero de 2026, proporcionando casi dos décadas de datos que incluyen múltiples "
        "ciclos económicos, tres crisis financieras importantes y diversos regímenes de mercado.")

    add_paragraph(doc,
        "La variable target para cada ETF se definió como el log-return de la semana siguiente "
        "(shift −1 sobre los retornos), lo que garantiza que las features contienen únicamente "
        "información del presente y el pasado mientras que el target es el retorno futuro. "
        "Esta separación temporal estricta es fundamental para evitar data leakage y asegurar "
        "que el modelo aprende a predecir el futuro utilizando solo información disponible en "
        "el momento de la decisión.")

    add_table(doc,
        headers=["Métrica", "Valor"],
        rows=[
            ["Filas (semanas)", "987"],
            ["Features predictoras", "109"],
            ["Variables target", "10 (una por ETF)"],
            ["Total columnas", "119"],
            ["Periodo", "2007-03-30 → 2026-02-20"],
            ["Nulos en features base", "0"],
            ["Nulos en features Refinitiv", "955 semanas (esperado: datos desde dic 2024)"],
            ["Archivos generados", "master_weekly.csv (normalizado), master_weekly_raw.csv"],
        ],
        col_widths=[5, 10.5])

    add_paragraph(doc,
        "El dataset se guarda en dos versiones: una normalizada (z-score sobre features) para "
        "exploración, y una sin normalizar que será la versión definitiva para el modelado, "
        "donde el scaler se ajustará exclusivamente con datos de entrenamiento.",
        space_before=6)


# ── Main ───────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generando documento Word...")

    doc = Document()

    # Configurar márgenes y pie de página
    setup_document(doc)

    # Portada (página sola con salto)
    crear_portada(doc)
    print("  Portada creada")

    # Título de sección principal
    add_heading_styled(doc, "3.2 Ingeniería del Dato", level=1)
    print("  Sección 3.2 iniciada")

    # Subsecciones
    seccion_321_origen(doc)
    print("  3.2.1 Origen de los datos")

    seccion_322_etl(doc)
    print("  3.2.2 Proceso ETL")

    seccion_323_limpieza(doc)
    print("  3.2.3 Limpieza y tratamiento")

    seccion_324_features(doc)
    print("  3.2.4 Feature Engineering")

    seccion_325_eda(doc)
    print("  3.2.5 Análisis Exploratorio")

    seccion_326_master(doc)
    print("  3.2.6 Dataset maestro")

    # Guardar
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    size_mb = os.path.getsize(OUTPUT) / (1024 * 1024)
    print(f"\nDocumento guardado: {OUTPUT} ({size_mb:.1f} MB)")
    print(f"Figuras insertadas: {fig_counter[0]}")
    print("Generación completada.")
